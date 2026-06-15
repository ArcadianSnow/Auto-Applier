"""Cover-letter autogen for strong matches — BUILD 5 (spec §6b, "ready just in case").

The optimize→apply pipeline already writes a per-job cover letter (the ``.txt`` the apply
driver pastes into a textarea). This module serves the **discovery+scoring-only** operating
mode, where the user applies externally and just wants a tailored letter sitting ready for
every strong job. It:

  1. generates the same guarded cover-letter body (:class:`CoverLetterGenerator` + the
     ``gen-cover-v4`` no-AI-tells prompt + the :func:`vet_cover_letter` fabrication guard), and
  2. renders it to a real Word **.docx** at the per-job upload path
     (``uploads/<job_id>/[<Name> ]Cover Letter.docx``) — ``.docx`` because the user pastes /
     uploads into web fields where markdown wouldn't render ([[feedback_paste_docs_as_docx]]),
     and because that path is exactly where the apply worker / ``av3 cover`` already look.

Invariants (mirror the apply path's reliability floor):

  * **NEVER clobber an existing letter.** A hand-authored ``av3 cover`` (or a prior autogen)
    always wins; autogen only fills the gap. A single deliberate regenerate may pass
    ``force=True`` (the ``av3 cover --generate <id> --force`` path); the batch backfill never
    forces.
  * **Guard fail-closed.** If :func:`vet_cover_letter` flags an unsupported tech claim the
    letter is NOT written and the job is left letterless with a note — a letter "ready just in
    case" must never be a fabrication the user would have to walk back live.
  * **No pipeline coupling.** This is CLI / daily-refresh driven (``av3 cover --generate`` /
    ``--generate-all``); it never advances job state and the apply worker never calls it.
"""

from __future__ import annotations

import re
import sqlite3
from dataclasses import dataclass
from pathlib import Path

from auto_applier.config.settings import Settings
from auto_applier.db.repositories import JobRepo, ScoreRepo
from auto_applier.domain.models import Job
from auto_applier.domain.state import JobState
from auto_applier.llm.complete import CompletionClient
from auto_applier.resume.factbank import Contact, FactBank
from auto_applier.resume.generate import (
    CoverLetterGenerator,
    existing_job_cover,
    job_cover_upload_path,
)
from auto_applier.resume.guard import vet_cover_letter

__all__ = [
    "CoverAutogenResult",
    "GENERATED",
    "SKIPPED_EXISTING",
    "SKIPPED_GUARD",
    "SKIPPED_NO_DESCRIPTION",
    "SKIPPED_DEGENERATE",
    "SKIPPED_INVALID",
    "ERROR",
    "backfill",
    "generate_one",
    "render_cover_letter_docx",
]

# --- result statuses --------------------------------------------------------
GENERATED = "generated"                      # wrote a fresh .docx
SKIPPED_EXISTING = "skipped_existing"        # a cover already exists (no-clobber); LLM not called
SKIPPED_GUARD = "skipped_guard"              # fabrication guard flagged unsupported claims; not written
SKIPPED_NO_DESCRIPTION = "skipped_no_description"  # no JD text to tailor against; LLM not called
SKIPPED_DEGENERATE = "skipped_degenerate"    # model produced runaway/repetitive output; not written
SKIPPED_INVALID = "skipped_invalid"          # model returned a non-letter (meta / recruiter-reply hallucination); not written
ERROR = "error"                              # generation or render raised


# Em-dash (U+2014) / en-dash (U+2013) used as a pause, with optional surrounding spaces but
# NOT across a newline (so paragraph breaks survive). The user's #1 AI tell.
_DASH_AS_PAUSE = re.compile(r"[ \t]*[—–][ \t]*")


def _strip_ai_tells(body: str) -> str:
    """Deterministic backstop for the #1 AI tell: em/en dashes.

    The ``gen-cover-v4`` prompt forbids them, but a local model drifts — so we also strip
    them mechanically here, replacing a dash-as-pause with a comma. This GUARANTEES no dash
    ever ships regardless of model behavior (the user is adamant). It only touches dashes;
    the rest of the no-AI-tells voice (excited/buzzwords/rule-of-three) stays prompt-driven,
    since those can't be fixed by a blind substitution without mangling meaning."""
    s = _DASH_AS_PAUSE.sub(", ", body or "")
    s = re.sub(r"[ \t]{2,}", " ", s)            # collapse space runs the substitution may leave
    s = re.sub(r"[ \t]+([,.;:])", r"\1", s)     # no space before punctuation
    s = re.sub(r",\s*,", ",", s)                # collapse a doubled comma
    return s


def _opens_in_third_person(body: str, name: str) -> bool:
    """True if the letter opens by naming the candidate or using 'He/His' as the subject.

    qwen3:8b occasionally drifts into third person on far-from-bank roles (observed on the
    Cockroach 'Value Engineer' JD, 2026-06-15: "Joseph Lira has built..."). A first-person
    letter always opens with 'I' or 'At/When/After <company>, I', so a name/He/His opener is
    an unambiguous, cheaply-detectable defect — :func:`generate_one` regenerates once on it."""
    head = (body or "").lstrip()[:80].lower()
    if not head:
        return False
    candidates = ["he ", "his "]
    n = (name or "").strip().lower()
    if n:
        candidates.append(n)            # full name as the opening subject
        candidates.append(n.split()[0] + " ")  # just the first name
    return any(head.startswith(c) for c in candidates)


_SENTENCE_SPLIT = re.compile(r"(?<=[.!?])\s+")


def _ensure_paragraphs(body: str) -> str:
    """Guarantee the hook / body / close 3-paragraph shape when the model returned one block.

    qwen3 honors the prompt's 'exactly three paragraphs' instruction only ~1/3 of the time
    (observed 2026-06-15). So if the body has no blank-line break and enough sentences,
    regroup deterministically: first sentence = the hook, last = the close, the rest = the
    middle. Content-preserving — it only inserts paragraph breaks at sentence boundaries, and
    leaves a body the model already paragraphed (or one too short to split) untouched."""
    text = (body or "").strip()
    if not text or "\n\n" in text:
        return text  # already paragraphed (or empty) — respect what the model produced
    sentences = [s.strip() for s in _SENTENCE_SPLIT.split(text) if s.strip()]
    if len(sentences) < 4:
        return text  # too short for a meaningful hook/body/close
    hook, close = sentences[0], sentences[-1]
    middle = " ".join(sentences[1:-1])
    return f"{hook}\n\n{middle}\n\n{close}"


# A sentence-initial standalone "I" — matches "I", "I've", "I built", "I also..."; does NOT
# match "In"/"It"/"If"/"After"/"At" (the word boundary fails on "In", which is a *good*,
# varied opener). Sentences are capitalized after the split, so a case-sensitive "I" is right.
_I_OPENER = re.compile(r"I\b")


def _i_sentence_flags(body: str) -> list[bool]:
    sentences = [s.strip() for s in _SENTENCE_SPLIT.split((body or "").strip()) if s.strip()]
    return [bool(_I_OPENER.match(s)) for s in sentences]


def _i_opening_ratio(body: str) -> float:
    """Fraction of sentences that begin with the word 'I'. 0.0 for an empty body — used as the
    monotony score so :func:`generate_one` can prefer the less-I-heavy of two drafts."""
    flags = _i_sentence_flags(body)
    return (sum(flags) / len(flags)) if flags else 0.0


def _excessive_i_openings(body: str) -> bool:
    """True if the draft is the monotone 'I did X. I did Y. I did Z.' résumé-in-prose the user
    rejected (2026-06-15: "it just keeps saying I did this I did this"): over 60% of a
    multi-sentence letter opens with 'I', OR three-plus 'I' openers in an unbroken run. The
    gen-cover-v4 prompt caps I-openers ("at most one sentence may begin with 'I'"), but soft
    style rules don't hold with qwen3 at greedy temp 0 — same lesson as third-person and
    degeneracy — so this gives the rule deterministic teeth. Unlike the degeneracy gate it
    never *skips* a letter (a monotone letter is still honest and usable); it only nudges
    generate_one toward the better of its two drafts."""
    flags = _i_sentence_flags(body)
    n = len(flags)
    if n < 3:
        return False  # too short to read as a "list"
    if n >= 4 and (sum(flags) / n) > 0.6:  # most of a multi-sentence letter opens with "I"
        return True
    run = best = 0
    for f in flags:
        run = run + 1 if f else 0
        best = max(best, run)
    return best >= 3  # three "I" openers in a row reads as a list even in a short letter


# qwen3:8b occasionally returns, as the "letter body", a META response — it role-plays the
# recruiter/evaluator instead of writing the candidate's letter. Observed live 2026-06-15 on the
# Databricks "Database Engine Internals" JD: "Joseph, your resume is impressive, but we need to
# see more about your experience with distributed systems... Could you please provide additional
# details?". This PASSES every other guard (no fabricated tech, short, opens "Joseph," not
# "He/His") and would SHIP. A real cover letter is written BY the candidate (first person about
# his own work), never addressed TO him by name, never asks him questions, never refers to "your
# resume". Detection is narrow to avoid false-flagging a normal "your team / your platform"
# (which correctly refers to the COMPANY).
_META_MARKERS = (
    "your resume", "your résumé", "your application", "your candidacy", "your background",
    "provide additional details", "provide more details", "we need to see",
    "we would like to see", "we'd like to see", "could you please provide",
    "tell us more about your", "more about your experience",
)


def _is_meta_response(body: str, name: str) -> bool:
    """True if the body is a recruiter/evaluator meta-response or otherwise not a letter the
    candidate wrote about himself — addresses him by name as a greeting, or uses evaluator
    phrasing aimed at him ('your resume', 'provide additional details')."""
    text = (body or "").strip()
    if not text:
        return False
    low = text.lower()
    head = low[:48]
    n = (name or "").strip().lower()
    greet_forms: list[str] = []
    if n:
        greet_forms += [n + ",", n + ":", n.split()[0] + ",", n.split()[0] + ":"]
    if any(head.startswith(g) for g in greet_forms):  # "Joseph," / "Joseph Lira:" — a letter is never TO him
        return True
    return any(m in low for m in _META_MARKERS)


# The prompt targets 150-250 words; 250 words is ~1700 chars, and a live batch of 458 letters
# put p95 at 1669 and p99 at 3101. So a body over ~2000 chars is over-target — either a qwen3
# repetition loop (the Mistral JDs ran to 20K-208K chars) or padded/doubled output. 2000 cleanly
# separates the 94% normal (<1500) tail from the runaway/padded ones without clipping a real letter.
_MAX_COVER_CHARS = 2000
# The clean samples have 5-7 sentences; a body with 13+ substantial sentences is a list-dump
# ("I built... I designed... I implemented..." ×24), not a letter — qwen3's other failure mode on
# far-from-bank ML/FD roles (unique sentences, so the repeat check misses it). 12 is ~2x the
# largest legitimate letter, so it never clips a real one.
_MAX_COVER_SENTENCES = 12


def _is_degenerate(body: str) -> bool:
    """True if the model produced runaway / padded / repetitive / list-dump output. qwen3 +
    greedy (temp 0) decoding can loop, repeating a sentence many times (Mistral JDs, 2026-06-15:
    20K-208K chars; milder doubled-sentence padding to ~3K chars), OR dump a wall of 24 distinct
    'I've done X' sentences. The fabrication guard passes all of these (every claim is
    bank-supported) and the dash/paragraph backstops don't check length, so this is the only
    thing standing between a degenerate generation and a shipped monstrosity. Caught by
    over-length OR too many sentences OR a substantial sentence repeated."""
    text = (body or "").strip()
    if len(text) > _MAX_COVER_CHARS:
        return True
    sentences = [p.strip() for p in _SENTENCE_SPLIT.split(text) if len(p.strip()) > 20]
    if len(sentences) > _MAX_COVER_SENTENCES:   # a wall of 13+ sentences is a list-dump, not a letter
        return True
    seen: dict[str, int] = {}
    for s in sentences:
        seen[s] = seen.get(s, 0) + 1
        if seen[s] >= 2:               # the same real sentence twice → padded/looping
            return True
    return False


@dataclass
class CoverAutogenResult:
    """One job's autogen outcome — observable, not side-effect-only (mirrors the
    pipeline run-summary style so the CLI can tally without re-querying)."""

    job_id: str
    status: str
    detail: str = ""
    path: str = ""  # the written .docx, set only when status == GENERATED

    @property
    def ok(self) -> bool:
        return self.status == GENERATED


# --- .docx render -----------------------------------------------------------

def render_cover_letter_docx(
    body: str,
    contact: Contact,
    out_path: Path,
    *,
    greeting: str = "Dear Hiring Manager,",
    closing: str = "Sincerely,",
) -> Path:
    """Render a complete, uploadable ``.docx`` cover letter.

    The generator returns a salutation/signature-less body (the apply driver wraps those when
    pasting into a textarea). A standalone document the human uploads needs to be *complete*,
    so we wrap a standard greeting, the body paragraphs, and a closing with the applicant's
    name. A name + contact header makes it read as a real letter.

    Anti-detection nicety: python-docx stamps ``author = "python-docx"`` in the file's core
    properties by default — a tell. We overwrite it with the applicant's own name (or blank).
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    name = (contact.name or "").strip() if contact else ""

    if name:
        head = doc.add_paragraph()
        run = head.add_run(name)
        run.bold = True
        run.font.size = Pt(14)

    contact_bits = [
        b.strip()
        for b in ((contact.email if contact else ""),
                  (contact.phone if contact else ""),
                  (contact.location if contact else ""))
        if b and b.strip()
    ]
    if contact_bits:
        doc.add_paragraph(" | ".join(contact_bits))

    if name or contact_bits:
        doc.add_paragraph("")  # spacer before the salutation

    doc.add_paragraph(greeting)
    doc.add_paragraph("")

    for para in (p.strip() for p in (body or "").split("\n\n")):
        if para:
            doc.add_paragraph(para)

    doc.add_paragraph("")
    doc.add_paragraph(closing)
    if name:
        doc.add_paragraph(name)

    try:  # core-properties author is metadata; never fatal if the backend rejects it
        doc.core_properties.author = name
    except Exception:  # noqa: BLE001
        pass

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# --- single job -------------------------------------------------------------

async def generate_one(
    settings: Settings,
    job: Job,
    *,
    bank: FactBank,
    generator: CoverLetterGenerator,
    name: str = "",
    vocabulary: tuple[str, ...] | None = None,
    force: bool = False,
) -> CoverAutogenResult:
    """Generate + guard + render ONE job's cover letter. Pure of state transitions.

    Order matches the invariants: no-clobber check (unless ``force``) → JD presence →
    LLM generation → fabrication guard (fail-closed) → ``.docx`` render. Any exception is
    caught and returned as an :data:`ERROR` result so a batch never aborts on one job.
    """
    if not force:
        existing = existing_job_cover(settings, job.id)
        if existing is not None:
            return CoverAutogenResult(
                job.id, SKIPPED_EXISTING,
                f"cover already present ({existing.name}); use --force to overwrite",
                str(existing),
            )

    if not (job.description or "").strip():
        return CoverAutogenResult(
            job.id, SKIPPED_NO_DESCRIPTION, "job has no description to tailor against"
        )

    # Generate, with one fail-safe second attempt. The retry serves two cases: (a) the first
    # draft opened in the third person (a clear defect the voice prompt mostly but not always
    # prevents), and (b) the first attempt errored — most often a JD whose qwen3 reasoning blew
    # past the Ollama read timeout. The retry drops qwen3's thinking (`no_think=True`), which
    # both rescues those slow JDs (~2s vs a 180s timeout) and regenerates a fresh first-person
    # draft. The default first attempt keeps thinking on, matching the bulk of letters.
    voice_name = (bank.contact.name if (bank and bank.contact) else "") or name
    # Best-of-two. A draft can be (a) degenerate (never shippable), (b) third-person or a
    # monotone "I did X. I did Y." list (shippable but a voice defect — worth a retry), or
    # (c) clean. Keep the BEST draft across attempts (clean > monotone/third-person >
    # degenerate) rather than blindly the last one, so the no_think retry can only improve the
    # result, never replace a good draft with a worse one. The retry also rescues a first
    # attempt that errored (a JD whose reasoning blew past the Ollama read timeout) and drops
    # qwen3 thinking (~2s vs a 180s timeout).
    best_body = ""
    best_key: tuple[int, int, float] | None = None  # (unshippable, third_person, i_ratio); lower=better
    last_exc: Exception | None = None
    for attempt in range(2):
        try:
            raw = await generator.generate(
                bank=bank,
                job_description=job.description,
                company=job.company,
                title=job.title,
                no_think=attempt > 0,
            )
        except ValueError:
            # parse rejected the reply: empty body OR non-letter JSON. qwen3 sometimes returns
            # an error-shaped / empty object for an odd JD (e.g. a Grafana-Tempo posting thick
            # with Prometheus config where "job" is a required label, 2026-06-15). That is a
            # CONTENT non-letter, not a transport crash — fall through to a clean skip, not ERROR.
            continue
        except Exception as exc:  # noqa: BLE001 — a real LLM/transport failure (timeout, conn)
            last_exc = exc
            continue  # retry once (the no_think attempt is fast and usually succeeds)
        last_exc = None
        cand = _strip_ai_tells(raw)  # deterministic em/en-dash backstop before scoring
        # `unshippable` = degenerate OR a meta/non-letter response — both can never ship, so they
        # rank below any real draft; a monotone or third-person draft is shippable but defective.
        unshippable = 1 if (_is_degenerate(cand) or _is_meta_response(cand, voice_name)) else 0
        key = (
            unshippable,
            1 if _opens_in_third_person(cand, voice_name) else 0,
            _i_opening_ratio(cand),
        )
        if best_key is None or key < best_key:
            best_body, best_key = cand, key
        # A clean draft (shippable, first person, not a monotone I-list) ends the loop; otherwise
        # fall through to the no_think retry, which may produce a less-defective draft.
        if key[0] == 0 and key[1] == 0 and not _excessive_i_openings(cand):
            break
    body = best_body
    if not body:
        if last_exc is not None:               # a real transport/LLM failure → ERROR (exit 1)
            return CoverAutogenResult(job.id, ERROR, f"generation failed: {last_exc}")
        # both attempts returned a non-letter (empty / error-shaped JSON) → clean skip, not a crash
        return CoverAutogenResult(
            job.id, SKIPPED_INVALID,
            "model returned no usable letter body (empty / non-letter JSON) after 2 attempts; not written",
        )
    if _is_degenerate(body):
        return CoverAutogenResult(
            job.id, SKIPPED_DEGENERATE,
            f"model produced runaway/repetitive output ({len(body)} chars); letter not written",
        )
    if _is_meta_response(body, voice_name):
        return CoverAutogenResult(
            job.id, SKIPPED_INVALID,
            "model returned a meta/recruiter-style response, not a cover letter; not written",
        )
    body = _ensure_paragraphs(body)  # guarantee hook/body/close shape if the model ran it together
    guard = vet_cover_letter(body, bank, vocabulary)
    if not guard.ok:
        terms = ", ".join(sorted({f.claim for f in guard.findings})) or "unsupported claim"
        return CoverAutogenResult(
            job.id, SKIPPED_GUARD,
            f"fabrication guard flagged: {terms} (letter not written)",
        )

    # force=True: clear every prior cover variant first so there's exactly one assigned file
    # (mirrors assign_cover_letter's replace semantics; uses only the public lookup).
    if force:
        for _ in range(8):  # bounded — at most a handful of ext variants ever exist
            prior = existing_job_cover(settings, job.id)
            if prior is None:
                break
            try:
                prior.unlink()
            except OSError:
                break

    out_path = job_cover_upload_path(settings, job.id, ".docx", name)
    try:
        render_cover_letter_docx(body, bank.contact, out_path)
    except Exception as exc:  # noqa: BLE001
        return CoverAutogenResult(job.id, ERROR, f"docx render failed: {exc}")

    return CoverAutogenResult(job.id, GENERATED, f"{job.company} — {job.title}", str(out_path))


# --- batch backfill ---------------------------------------------------------

async def backfill(
    settings: Settings,
    conn: sqlite3.Connection,
    *,
    llm: CompletionClient,
    bank: FactBank,
    min_score: float,
    name: str = "",
    limit: int | None = None,
    states: tuple[JobState, ...] = (JobState.DECIDED,),
    vocabulary: tuple[str, ...] | None = None,
) -> list[CoverAutogenResult]:
    """Write letters for every job scoring ≥ ``min_score`` (in ``states``) that lacks one.

    Reads :meth:`ScoreRepo.list_ranked` (ranked by score desc) with ``min_total=min_score``,
    filters to ``states`` (default ``DECIDED`` — a strong match resting in discovery+scoring-
    only mode), skips jobs that already have a cover (recorded, but they don't consume
    ``limit``), and generates up to ``limit`` *new* letters. The backfill never forces, so a
    hand-authored letter is always preserved. Returns one result per job touched, highest
    score first — the CLI tallies them.
    """
    generator = CoverLetterGenerator(llm)
    job_repo = JobRepo(conn)
    allowed = {s.value for s in states}

    results: list[CoverAutogenResult] = []
    generated = 0  # counts only jobs where the LLM actually ran (caps real work)

    for row in ScoreRepo(conn).list_ranked(min_total=min_score):
        if row.get("state") not in allowed:
            continue
        job = job_repo.get(row["job_id"])
        if job is None:
            continue

        # An existing letter is reported but is free (no LLM) — don't let it eat the limit.
        if existing_job_cover(settings, job.id) is not None:
            results.append(CoverAutogenResult(
                job.id, SKIPPED_EXISTING, f"{job.company} — {job.title}"
            ))
            continue

        if limit is not None and generated >= limit:
            break

        res = await generate_one(
            settings, job, bank=bank, generator=generator, name=name, vocabulary=vocabulary
        )
        results.append(res)
        # GENERATED / SKIPPED_GUARD / SKIPPED_DEGENERATE / SKIPPED_INVALID / ERROR all mean the
        # LLM was invoked → consume a slot. SKIPPED_NO_DESCRIPTION short-circuits before the LLM.
        if res.status in (GENERATED, SKIPPED_GUARD, SKIPPED_DEGENERATE, SKIPPED_INVALID, ERROR):
            generated += 1

    return results
