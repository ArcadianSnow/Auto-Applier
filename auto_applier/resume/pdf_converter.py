"""Resume format normalization to PDF.

Job sites overwhelmingly prefer PDF and many outright reject .doc /
.docx (Workday, some Greenhouse forms). The user keeps their original
resume in whatever format they already had — we never touch it — but
the upload layer always sends a PDF derived from that file.

Public surface: :func:`ensure_pdf`. Cached output lives under
``data/resumes/.converted/<stem>.pdf`` and is regenerated when the
source file's mtime advances past the cache file's mtime.
"""
from __future__ import annotations

import asyncio
import html
import logging
import re
from pathlib import Path

from auto_applier.config import CONVERTED_RESUMES_DIR

logger = logging.getLogger(__name__)


async def ensure_pdf(resume_path: str | Path) -> Path:
    """Return a PDF path for ``resume_path``, converting if necessary.

    - ``.pdf`` files are returned unchanged.
    - ``.docx`` and ``.txt`` files are rendered to PDF via the same
      Playwright HTML→PDF pipeline used by tailored resumes and cover
      letters. Cached under ``data/resumes/.converted/`` and reused
      until the source file is newer than the cached output.
    - ``.doc`` (legacy binary) and other formats fall back to the
      original path with a warning — converting binary .doc reliably
      requires LibreOffice or Word, which violates the no-cost rule.

    Raises FileNotFoundError if the source doesn't exist. Returns the
    original path on any conversion failure (so the caller can still
    attempt the upload — graceful degradation beats hard-fail).
    """
    src = Path(resume_path)
    if not src.exists():
        raise FileNotFoundError(f"Resume not found: {src}")

    suffix = src.suffix.lower()
    if suffix == ".pdf":
        # Always make a working copy under .converted/ even when input
        # is already PDF. Two reasons:
        #   1. Read/write code paths stay uniform — every consumer
        #      (text extraction for scoring, set_input_files for
        #      upload) reads from .converted/ and never touches the
        #      original file. So if the user opens the original in
        #      Adobe Acrobat, that lock can't disrupt a run.
        #   2. The cache is mtime-keyed, so editing the source
        #      replaces the cache automatically.
        cached = _cached_pdf_path(src)
        if _cache_is_fresh(src, cached):
            return cached
        try:
            import shutil
            CONVERTED_RESUMES_DIR.mkdir(parents=True, exist_ok=True)
            shutil.copy2(src, cached)
            logger.info(
                "Cached PDF resume: %s → %s", src.name, cached,
            )
            return cached
        except Exception as exc:
            logger.warning(
                "Failed to cache PDF resume %s (%s) — uploading "
                "original. Keep the file out of Acrobat to avoid "
                "lock issues.", src, exc,
            )
            return src

    if suffix not in (".docx", ".txt"):
        logger.warning(
            "Resume format %s not auto-convertible to PDF — uploading "
            "as-is. Consider exporting %s to PDF for better acceptance "
            "rates on Indeed / LinkedIn / Workday.",
            suffix, src.name,
        )
        return src

    cached = _cached_pdf_path(src)
    if _cache_is_fresh(src, cached):
        return cached

    try:
        if suffix == ".docx":
            html_body = _docx_to_html(src)
        else:  # .txt
            html_body = _text_to_html(src.read_text(encoding="utf-8"))
    except Exception as exc:
        logger.warning(
            "Failed to read %s for PDF conversion: %s — uploading "
            "original file as-is.", src, exc,
        )
        return src

    if not html_body.strip():
        logger.warning(
            "Resume %s produced empty HTML after parsing — uploading "
            "original.", src,
        )
        return src

    full_html = _wrap_resume_html(html_body)
    ok = await _render(full_html, cached)
    if ok:
        logger.info("Converted resume to PDF: %s → %s", src.name, cached)
        return cached
    logger.warning(
        "PDF render failed for %s — uploading original.", src,
    )
    return src


# ---------------------------------------------------------------------------
# Cache helpers
# ---------------------------------------------------------------------------


def _cached_pdf_path(src: Path) -> Path:
    CONVERTED_RESUMES_DIR.mkdir(parents=True, exist_ok=True)
    safe_stem = re.sub(r"[^a-zA-Z0-9._-]", "_", src.stem)
    return CONVERTED_RESUMES_DIR / f"{safe_stem}.pdf"


def _cache_is_fresh(src: Path, cached: Path) -> bool:
    """True if cached PDF exists and is newer than the source."""
    if not cached.exists():
        return False
    try:
        return cached.stat().st_mtime >= src.stat().st_mtime
    except OSError:
        return False


# ---------------------------------------------------------------------------
# DOCX → HTML
# ---------------------------------------------------------------------------


def _docx_to_html(path: Path) -> str:
    """Convert a .docx to a minimal HTML body.

    Uses python-docx (already a dependency for resume parsing). Emits
    paragraphs, headings, and bullet lists. Preserves runs of bold /
    italic text. Does not preserve tables or images — keeping the
    converter dependency-free; tables become flattened paragraphs.
    """
    from docx import Document

    with open(path, "rb") as fh:
        doc = Document(fh)

    parts: list[str] = []
    in_list = False
    for para in doc.paragraphs:
        text = para.text or ""
        style = (para.style.name or "").lower() if para.style else ""

        if not text.strip():
            if in_list:
                parts.append("</ul>")
                in_list = False
            parts.append("<p>&nbsp;</p>")
            continue

        # Bullet list?
        if "list" in style or "bullet" in style:
            if not in_list:
                parts.append("<ul>")
                in_list = True
            parts.append(f"<li>{_runs_to_html(para.runs, text)}</li>")
            continue
        if in_list:
            parts.append("</ul>")
            in_list = False

        # Headings
        if "heading 1" in style or "title" in style:
            parts.append(f"<h1>{html.escape(text)}</h1>")
            continue
        if "heading 2" in style:
            parts.append(f"<h2>{html.escape(text)}</h2>")
            continue
        if "heading" in style:
            parts.append(f"<h3>{html.escape(text)}</h3>")
            continue

        parts.append(f"<p>{_runs_to_html(para.runs, text)}</p>")

    if in_list:
        parts.append("</ul>")

    # Tables as flattened paragraphs (no styling — we just want the
    # text to survive into the PDF for ATS parsing).
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                (cell.text or "").strip() for cell in row.cells
            )
            row_text = row_text.strip(" |")
            if row_text:
                parts.append(f"<p>{html.escape(row_text)}</p>")

    return "\n".join(parts)


def _runs_to_html(runs, fallback_text: str) -> str:
    """Render a paragraph's runs preserving bold/italic. Falls back to
    escaping the whole paragraph text if the run list is empty."""
    if not runs:
        return html.escape(fallback_text)
    out: list[str] = []
    for run in runs:
        text = html.escape(run.text or "")
        if not text:
            continue
        if getattr(run, "bold", False):
            text = f"<strong>{text}</strong>"
        if getattr(run, "italic", False):
            text = f"<em>{text}</em>"
        out.append(text)
    return "".join(out) or html.escape(fallback_text)


# ---------------------------------------------------------------------------
# Plain text → HTML
# ---------------------------------------------------------------------------


def _text_to_html(text: str) -> str:
    """Convert plain text to a minimal HTML body.

    Heuristics:
      - Blank-line-separated blocks become <p> elements
      - Lines starting with ``-``, ``*``, or ``•`` become <li> in a <ul>
      - All-caps short lines become <h2> (typical resume section headers)
    """
    lines = text.replace("\r\n", "\n").split("\n")
    parts: list[str] = []
    buf: list[str] = []
    in_list = False

    def flush_buf() -> None:
        nonlocal buf
        if not buf:
            return
        joined = " ".join(b.strip() for b in buf if b.strip())
        if joined:
            parts.append(f"<p>{html.escape(joined)}</p>")
        buf = []

    def close_list() -> None:
        nonlocal in_list
        if in_list:
            parts.append("</ul>")
            in_list = False

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            flush_buf()
            close_list()
            continue
        # Bullet?
        m = re.match(r"^\s*[-*•]\s+(.*)$", line)
        if m:
            flush_buf()
            if not in_list:
                parts.append("<ul>")
                in_list = True
            parts.append(f"<li>{html.escape(m.group(1).strip())}</li>")
            continue
        close_list()
        # ALL-CAPS short line → heading
        stripped = line.strip()
        if (len(stripped) <= 60 and stripped.isupper()
                and any(c.isalpha() for c in stripped)):
            flush_buf()
            parts.append(f"<h2>{html.escape(stripped)}</h2>")
            continue
        buf.append(line)

    flush_buf()
    close_list()
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# HTML wrapper + PDF render
# ---------------------------------------------------------------------------


_RESUME_CSS = """
@page { size: Letter; margin: 0.5in; }
body { font-family: 'Calibri', 'Helvetica', sans-serif;
       font-size: 11pt; line-height: 1.35; color: #111; }
h1 { font-size: 18pt; margin: 0 0 4pt 0; }
h2 { font-size: 13pt; margin: 12pt 0 4pt 0;
     border-bottom: 1px solid #888; padding-bottom: 2pt; }
h3 { font-size: 11.5pt; margin: 8pt 0 2pt 0; }
p { margin: 2pt 0; }
ul { margin: 2pt 0 6pt 18pt; padding: 0; }
li { margin: 2pt 0; }
strong { font-weight: 600; }
em { font-style: italic; }
"""


def _wrap_resume_html(body: str) -> str:
    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        f"<style>{_RESUME_CSS}</style></head><body>"
        f"{body}</body></html>"
    )


async def _render(html_content: str, out_path: Path) -> bool:
    """Render via tailor.render_pdf (Playwright). Imported lazily so
    test environments without Playwright installed can still import
    this module."""
    try:
        from auto_applier.resume.tailor import render_pdf
    except Exception as exc:
        logger.warning("Cannot import render_pdf: %s", exc)
        return False
    try:
        return await render_pdf(html_content, out_path)
    except Exception as exc:
        logger.warning("render_pdf raised: %s", exc)
        return False


def ensure_pdf_sync(resume_path: str | Path) -> Path:
    """Synchronous wrapper around :func:`ensure_pdf`.

    For callers that aren't in an async context (CLI tools, doctor
    checks, the wizard validation step). Spins up an event loop only
    if there isn't one already running.
    """
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            # Caller is in an async context but used the sync wrapper —
            # they should await ensure_pdf directly. Return the source
            # path as a safe no-op rather than blocking.
            logger.warning(
                "ensure_pdf_sync called from inside a running event "
                "loop; returning original path. Use `await ensure_pdf` "
                "in async code."
            )
            return Path(resume_path)
    except RuntimeError:
        pass
    return asyncio.run(ensure_pdf(resume_path))
