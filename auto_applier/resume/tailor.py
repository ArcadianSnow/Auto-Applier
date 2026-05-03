"""Per-JD tailored resume generation.

Given a stored job and a base resume, produce a rewritten version
that emphasizes the material most relevant to that specific listing,
render it through a simple HTML template, and save the result as a
PDF under ``data/profiles/generated/<job_id>.pdf``.

The original resume files in ``data/resumes/`` are **never**
modified — that constraint is load-bearing for the whole project.
Tailored PDFs are generated artifacts, not edits.

PDF rendering uses the same Playwright Chromium we already require
for scraping. No new dependencies. The FormFiller's resume-upload
path should prefer a tailored version when one exists for the
target job_id.
"""

from __future__ import annotations

import html
import json
import logging
from dataclasses import dataclass
from pathlib import Path

from auto_applier.config import GENERATED_RESUMES_DIR
from auto_applier.llm.prompts import TAILOR_RESUME
from auto_applier.llm.router import LLMRouter

logger = logging.getLogger(__name__)


@dataclass
class TailoredResume:
    summary: str
    skills: list[str]
    experience: list[dict]  # each: {title, company, dates, bullets}
    education: list[dict]   # each: {school, degree, year}
    job_id: str = ""
    source_resume_label: str = ""


# Minimal ATS-friendly HTML template. No fancy CSS, no web fonts,
# no images — just clean semantic structure that parses well in
# resume trackers and prints to a readable single column.
_HTML_TEMPLATE = """<!DOCTYPE html>
<html><head>
<meta charset="utf-8">
<title>Resume</title>
<style>
  body {{ font-family: Arial, Helvetica, sans-serif; font-size: 11pt;
          color: #111; max-width: 720px; margin: 24px auto; line-height: 1.45; }}
  h1 {{ font-size: 22pt; margin: 0 0 4pt 0; }}
  h2 {{ font-size: 13pt; margin: 18pt 0 4pt 0; border-bottom: 1px solid #888;
        padding-bottom: 2pt; text-transform: uppercase; letter-spacing: 0.5pt; }}
  h3 {{ font-size: 11pt; margin: 10pt 0 2pt 0; }}
  .role-meta {{ color: #555; font-size: 10pt; margin: 0 0 4pt 0; }}
  ul {{ margin: 4pt 0 8pt 0; padding-left: 18pt; }}
  li {{ margin: 2pt 0; }}
  .skills {{ margin: 4pt 0; }}
  .header-contact {{ color: #333; font-size: 10pt; margin: 0 0 12pt 0; }}
</style>
</head>
<body>
<h1>{name}</h1>
<div class="header-contact">{contact}</div>
<h2>Summary</h2>
<p>{summary}</p>
<h2>Skills</h2>
<div class="skills">{skills}</div>
<h2>Experience</h2>
{experience}
<h2>Education</h2>
{education}
</body></html>
"""


def render_html(tailored: TailoredResume, name: str = "", contact: str = "") -> str:
    """Render a TailoredResume to self-contained HTML."""
    def _esc(s: str) -> str:
        return html.escape(str(s or ""))

    skills_html = ", ".join(_esc(s) for s in tailored.skills)
    experience_blocks = []
    for job in tailored.experience:
        bullets = "".join(f"<li>{_esc(b)}</li>" for b in job.get("bullets", []))
        experience_blocks.append(
            f'<h3>{_esc(job.get("title", ""))} — {_esc(job.get("company", ""))}</h3>'
            f'<p class="role-meta">{_esc(job.get("dates", ""))}</p>'
            f'<ul>{bullets}</ul>'
        )
    experience_html = "\n".join(experience_blocks) or "<p><em>No experience listed.</em></p>"

    education_blocks = []
    for edu in tailored.education:
        education_blocks.append(
            f'<p><strong>{_esc(edu.get("school", ""))}</strong> — '
            f'{_esc(edu.get("degree", ""))} '
            f'<span class="role-meta">({_esc(edu.get("year", ""))})</span></p>'
        )
    education_html = "\n".join(education_blocks) or "<p><em>No education listed.</em></p>"

    return _HTML_TEMPLATE.format(
        name=_esc(name or "Candidate"),
        contact=_esc(contact),
        summary=_esc(tailored.summary),
        skills=skills_html,
        experience=experience_html,
        education=education_html,
    )


class ResumeTailor:
    """Rewrites a resume for a specific job description via the LLM router."""

    def __init__(self, router: LLMRouter) -> None:
        self.router = router

    async def tailor(
        self,
        resume_text: str,
        job_description: str,
        company_name: str,
        job_title: str,
        job_id: str = "",
        resume_label: str = "",
    ) -> TailoredResume | None:
        """Produce a TailoredResume, or None on LLM failure.

        Validates that the response has at least a summary and some
        skills — a shell response with empty fields would produce
        a useless PDF and is rejected.
        """
        try:
            result = await self.router.complete_json(
                prompt=TAILOR_RESUME.format(
                    resume_text=resume_text[:4000],
                    job_description=job_description[:2000],
                    company_name=company_name,
                    job_title=job_title,
                ),
                system_prompt=TAILOR_RESUME.system,
            )
        except Exception as e:
            logger.debug("Resume tailor LLM call failed: %s", e)
            return None

        summary = str(result.get("summary", "")).strip()
        skills = result.get("skills", [])
        if not summary or not isinstance(skills, list) or not skills:
            logger.debug("Tailor response too sparse, discarding")
            return None

        experience = result.get("experience", [])
        if not isinstance(experience, list):
            experience = []
        education = result.get("education", [])
        if not isinstance(education, list):
            education = []

        return TailoredResume(
            summary=summary,
            skills=[str(s) for s in skills if str(s).strip()],
            experience=[e for e in experience if isinstance(e, dict)],
            education=[e for e in education if isinstance(e, dict)],
            job_id=job_id,
            source_resume_label=resume_label,
        )


async def render_pdf(html_content: str, out_path: Path) -> bool:
    """Render HTML to a PDF at ``out_path`` via Playwright Chromium.

    Returns True on success, False on any failure. The browser
    context is ephemeral — we don't reuse the main scraping context
    because PDF rendering happens outside a platform flow.
    """
    try:
        from playwright.async_api import async_playwright
    except ImportError:
        logger.warning("Playwright not installed — cannot render PDF")
        return False

    try:
        async with async_playwright() as pw:
            browser = await pw.chromium.launch(headless=True)
            context = await browser.new_context()
            page = await context.new_page()
            await page.set_content(html_content, wait_until="domcontentloaded")
            out_path.parent.mkdir(parents=True, exist_ok=True)
            await page.pdf(
                path=str(out_path),
                format="Letter",
                margin={"top": "0.5in", "bottom": "0.5in",
                        "left": "0.5in", "right": "0.5in"},
                print_background=True,
            )
            await context.close()
            await browser.close()
        return True
    except Exception as e:
        logger.warning("PDF render failed: %s", e)
        return False


def _user_filename_prefix() -> str:
    """Build a 'First_Last' or 'First' filename prefix from user_config.

    The previous naming pattern was `<job_id>.pdf` (e.g.
    `ind-9934dbc8cae647b8.pdf`) — uploading a file with that name
    is a dead giveaway that the resume was system-generated. Using
    the candidate's actual name reads as if they exported their own
    resume from Word.

    Falls back to "Resume" if user_config.json is missing or has no
    name fields, so the filename is still organic-looking.
    """
    try:
        import json
        from auto_applier.config import USER_CONFIG_FILE
        if USER_CONFIG_FILE.exists():
            data = json.loads(USER_CONFIG_FILE.read_text(encoding="utf-8"))
            personal = data.get("personal_info", {}) or {}
            first = (personal.get("first_name") or "").strip()
            last = (personal.get("last_name") or "").strip()
            if first and last:
                return f"{_clean_name(first)}_{_clean_name(last)}"
            if first:
                return _clean_name(first)
            full = (personal.get("name") or "").strip()
            if full:
                return _clean_name(full).replace(" ", "_")
    except Exception:
        pass
    return ""


def _clean_name(name: str) -> str:
    """Strip filesystem-unfriendly chars from a name."""
    return "".join(c for c in name if c.isalnum() or c in "_- ").strip()


def tailored_pdf_path(job_id: str) -> Path:
    """Return the canonical path for a tailored PDF for a job.

    Layout: ``<GENERATED_RESUMES_DIR>/<job_id>/<First>_<Last>_Resume.pdf``.
    The job_id stays in the path (we need a stable lookup key) but
    sits in the *directory* layer, not the filename — so the basename
    that gets uploaded to a job board is ``Jordan_Testpilot_Resume.pdf``,
    not ``ind-9934dbc8cae647b8.pdf``.
    """
    # Sanitize job_id for filesystem safety
    safe_job = "".join(c if c.isalnum() or c in "-_." else "_" for c in job_id)
    prefix = _user_filename_prefix()
    filename = f"{prefix}_Resume.pdf" if prefix else "Resume.pdf"
    return GENERATED_RESUMES_DIR / safe_job / filename


def tailored_docx_path(job_id: str) -> Path:
    """Sibling of :func:`tailored_pdf_path` for the .docx variant.

    Workday and Taleo parse DOCX with ~97% field-extraction accuracy
    vs. ~83% for PDF (per the Phase 1 resume-tailoring research).
    Greenhouse/Lever/Ashby/Indeed/Dice/ZR all parse PDF fine, so PDF
    stays the default — but we generate DOCX alongside so the form
    filler can prefer it on platforms where it's empirically better.
    """
    safe_job = "".join(c if c.isalnum() or c in "-_." else "_" for c in job_id)
    prefix = _user_filename_prefix()
    filename = f"{prefix}_Resume.docx" if prefix else "Resume.docx"
    return GENERATED_RESUMES_DIR / safe_job / filename


def archetype_tailored_pdf_path(resume_label: str, archetype: str) -> Path:
    """Cache path for archetype-tailored resumes (Phase 2.3).

    The Phase 2 plan has three cache tiers:
      L1: per-job tailored at ``GENERATED_RESUMES_DIR/<job_id>/...``
          (set by Phase 1.6 ``_tailor_resume_for_job`` on first apply)
      L2: per-archetype tailored at
          ``GENERATED_RESUMES_DIR/_archetypes/<resume_label>/<archetype>/...``
          (set by `cli refresh-tailored-resumes` in idle time, or
          on-demand)
      L3: base resume from ``data/resumes/`` (always-available
          fallback)

    Layout uses ``_archetypes`` (leading underscore) as the parent
    so it sorts alphabetically away from per-job dirs and is easy
    to spot in `data/profiles/generated/`. Subdirs by resume_label
    so two resumes (e.g. ``Data_Analyst`` and ``Data_Engineer``)
    don't collide on shared archetype names.
    """
    safe_label = "".join(
        c if c.isalnum() or c in "-_." else "_"
        for c in (resume_label or "default")
    )
    safe_arch = "".join(
        c if c.isalnum() or c in "-_." else "_"
        for c in (archetype or "default")
    )
    prefix = _user_filename_prefix()
    filename = f"{prefix}_Resume.pdf" if prefix else "Resume.pdf"
    return (
        GENERATED_RESUMES_DIR / "_archetypes" / safe_label / safe_arch / filename
    )


def archetype_tailored_docx_path(resume_label: str, archetype: str) -> Path:
    """DOCX sibling of :func:`archetype_tailored_pdf_path`."""
    safe_label = "".join(
        c if c.isalnum() or c in "-_." else "_"
        for c in (resume_label or "default")
    )
    safe_arch = "".join(
        c if c.isalnum() or c in "-_." else "_"
        for c in (archetype or "default")
    )
    prefix = _user_filename_prefix()
    filename = f"{prefix}_Resume.docx" if prefix else "Resume.docx"
    return (
        GENERATED_RESUMES_DIR / "_archetypes" / safe_label / safe_arch / filename
    )


async def render_docx(tailored: TailoredResume, out_path: Path,
                      name: str = "", contact: str = "") -> bool:
    """Render a TailoredResume to a single-column ATS-friendly DOCX.

    Per Phase 1 research: Workday/Taleo prefer DOCX (97% extraction
    vs 83% for PDF), iCIMS slightly prefers DOCX too. Single-column,
    plain-font is mandatory for parser-friendliness — multi-column
    layouts and graphical skill bars still break parsers in 2026.

    Returns True on success, False on any failure (caller falls back
    to PDF). Async-signatured to mirror :func:`render_pdf` so the
    engine's call sites stay symmetric.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        logger.warning("python-docx not installed — cannot render DOCX")
        return False

    try:
        doc = Document()

        # Page margins — 0.6" matches the 0.5" PDF margin closely.
        for section in doc.sections:
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)

        # Set base font for the whole doc to a parser-friendly
        # plain font. ATS parsers fail on uncommon fonts.
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Header — name + contact info
        if name:
            heading = doc.add_paragraph()
            run = heading.add_run(name)
            run.bold = True
            run.font.size = Pt(20)
        if contact:
            doc.add_paragraph(contact)

        def _section(title: str) -> None:
            """Add a section heading; matches the PDF template's
            uppercased + small-letter-spaced look."""
            p = doc.add_paragraph()
            run = p.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(12)

        # Summary
        _section("Summary")
        doc.add_paragraph(tailored.summary or "")

        # Skills — comma-joined for easier parsing than
        # bullet-per-skill (parsers handle prose better than columns)
        _section("Skills")
        if tailored.skills:
            doc.add_paragraph(", ".join(str(s) for s in tailored.skills))

        # Experience
        _section("Experience")
        for exp in tailored.experience:
            if not isinstance(exp, dict):
                continue
            title_run = doc.add_paragraph()
            r = title_run.add_run(
                f"{exp.get('title', '')} — {exp.get('company', '')}"
            )
            r.bold = True
            dates = exp.get("dates", "")
            if dates:
                meta = doc.add_paragraph(dates)
                # Italic, smaller — same emphasis as PDF role-meta
                for run in meta.runs:
                    run.italic = True
                    run.font.size = Pt(10)
            for bullet in exp.get("bullets", []) or []:
                doc.add_paragraph(str(bullet), style="List Bullet")

        # Education
        _section("Education")
        for edu in tailored.education:
            if not isinstance(edu, dict):
                continue
            line = (
                f"{edu.get('school', '')} — "
                f"{edu.get('degree', '')} "
                f"({edu.get('year', '')})"
            ).strip()
            doc.add_paragraph(line)

        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        return True
    except Exception as exc:
        logger.warning("DOCX render failed: %s", exc)
        return False


def save_tailored_json(tailored: TailoredResume) -> Path:
    """Persist the structured TailoredResume next to the PDF."""
    from dataclasses import asdict

    path = tailored_pdf_path(tailored.job_id).with_suffix(".json")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(asdict(tailored), indent=2), encoding="utf-8",
    )
    return path
