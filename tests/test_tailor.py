"""Tests for per-JD resume tailoring — excludes live Playwright PDF render."""
import asyncio
from unittest.mock import AsyncMock, MagicMock

import pytest

from auto_applier.resume.tailor import (
    ResumeTailor,
    TailoredResume,
    render_docx,
    render_html,
    tailored_docx_path,
    tailored_pdf_path,
)


class TestTailoredPdfPath:
    """job_id sanitization moved from filename to directory: the
    PDF is named after the user (or 'Resume.pdf' fallback) so the
    upload's basename doesn't betray the system. The job_id lives
    in the parent directory."""

    def test_sanitizes_slashes_in_job_dir(self):
        path = tailored_pdf_path("li/../etc/passwd")
        # Slashes (the path-traversal vector) can't appear in the
        # job_id segment of the path. Bare ".." inside a single
        # segment is not a traversal — it's just two characters.
        assert "/" not in path.parent.name
        assert "\\" not in path.parent.name
        assert path.suffix == ".pdf"

    def test_preserves_safe_chars_in_job_dir(self):
        path = tailored_pdf_path("li-12345_abc.def")
        assert path.parent.name == "li-12345_abc.def"

    def test_strips_spaces_from_job_dir(self):
        path = tailored_pdf_path("job 1 2 3")
        assert " " not in path.parent.name

    def test_filename_does_not_contain_job_id(self):
        """Critical: uploading a file named after the job_id would
        be a dead giveaway that the resume is system-generated."""
        path = tailored_pdf_path("ind-9934dbc8cae647b8")
        assert "ind-9934" not in path.name
        assert path.name.endswith("_Resume.pdf") or path.name == "Resume.pdf"


class TestRenderHtml:
    def test_minimal_render(self):
        t = TailoredResume(
            summary="Senior analyst with 5 years of SQL experience.",
            skills=["Python", "SQL", "Tableau"],
            experience=[],
            education=[],
        )
        html = render_html(t, name="Jane Doe", contact="jane@x.com")
        assert "Jane Doe" in html
        assert "jane@x.com" in html
        assert "Senior analyst with 5 years" in html
        assert "Python, SQL, Tableau" in html

    def test_experience_block(self):
        t = TailoredResume(
            summary="x",
            skills=["y"],
            experience=[{
                "title": "Analyst",
                "company": "Acme",
                "dates": "2020-2023",
                "bullets": ["Built dashboards", "Led SQL training"],
            }],
            education=[],
        )
        html = render_html(t, name="Jane", contact="")
        assert "Analyst — Acme" in html
        assert "2020-2023" in html
        assert "Built dashboards" in html
        assert "Led SQL training" in html

    def test_education_block(self):
        t = TailoredResume(
            summary="x", skills=["y"], experience=[],
            education=[{"school": "State", "degree": "BS CS", "year": "2018"}],
        )
        html = render_html(t, "Jane", "")
        assert "State" in html
        assert "BS CS" in html
        assert "2018" in html

    def test_empty_sections_say_none(self):
        t = TailoredResume(summary="x", skills=["y"], experience=[], education=[])
        html = render_html(t, "Jane", "")
        assert "No experience listed" in html
        assert "No education listed" in html

    def test_html_escaping(self):
        t = TailoredResume(
            summary="Had <script>alert(1)</script> exp",
            skills=["a & b"],
            experience=[],
            education=[],
        )
        html = render_html(t, "<b>Jane</b>", "x@y.com")
        assert "<script>" not in html
        assert "&lt;script&gt;" in html
        assert "&lt;b&gt;Jane&lt;/b&gt;" in html
        assert "a &amp; b" in html


class TestResumeTailor:
    def _router_returning(self, payload: dict):
        router = MagicMock()
        router.complete_json = AsyncMock(return_value=payload)
        return router

    def test_returns_none_on_empty_summary(self):
        router = self._router_returning({
            "summary": "",
            "skills": ["Python"],
        })
        result = asyncio.run(ResumeTailor(router).tailor(
            resume_text="x", job_description="y",
            company_name="Acme", job_title="Analyst",
        ))
        assert result is None

    def test_returns_none_on_no_skills(self):
        router = self._router_returning({
            "summary": "Strong analyst",
            "skills": [],
        })
        result = asyncio.run(ResumeTailor(router).tailor(
            resume_text="x", job_description="y",
            company_name="Acme", job_title="Analyst",
        ))
        assert result is None

    def test_returns_none_on_llm_exception(self):
        router = MagicMock()
        router.complete_json = AsyncMock(side_effect=RuntimeError("boom"))
        result = asyncio.run(ResumeTailor(router).tailor(
            resume_text="x", job_description="y",
            company_name="Acme", job_title="Analyst",
        ))
        assert result is None

    def test_accepts_complete_response(self):
        router = self._router_returning({
            "summary": "Analyst with dashboards.",
            "skills": ["SQL", "Python"],
            "experience": [{
                "title": "Analyst", "company": "Acme",
                "dates": "2020-2023", "bullets": ["bullet a"],
            }],
            "education": [{"school": "State", "degree": "BS", "year": "2018"}],
        })
        result = asyncio.run(ResumeTailor(router).tailor(
            resume_text="x", job_description="y",
            company_name="Acme", job_title="Analyst",
            job_id="li-123", resume_label="analyst",
        ))
        assert result is not None
        assert result.summary.startswith("Analyst")
        assert result.skills == ["SQL", "Python"]
        assert len(result.experience) == 1
        assert result.education[0]["school"] == "State"
        assert result.job_id == "li-123"
        assert result.source_resume_label == "analyst"

    def test_filters_non_dict_experience(self):
        router = self._router_returning({
            "summary": "x",
            "skills": ["SQL"],
            "experience": [
                {"title": "Analyst", "company": "Acme", "dates": "y", "bullets": []},
                "not a dict",
                42,
            ],
            "education": [],
        })
        result = asyncio.run(ResumeTailor(router).tailor(
            resume_text="x", job_description="y",
            company_name="Acme", job_title="Analyst",
        ))
        assert len(result.experience) == 1

    def test_stringifies_skill_entries(self):
        router = self._router_returning({
            "summary": "x",
            "skills": ["SQL", 42, None, ""],
        })
        result = asyncio.run(ResumeTailor(router).tailor(
            resume_text="x", job_description="y",
            company_name="Acme", job_title="Analyst",
        ))
        # Non-empty stringified values kept, empty/None filtered
        assert "SQL" in result.skills
        assert "42" in result.skills
        assert "None" in result.skills  # None→"None" is non-empty
        assert "" not in result.skills


# ----------------------------------------------------------------------
# DOCX render — Phase 1 addition for Workday/Taleo parser preference
# ----------------------------------------------------------------------

class TestTailoredDocxPath:
    """DOCX path mirrors the PDF path layout: <job_dir>/<name>_Resume.docx
    so the basename uploaded to a job board is candidate-named, not
    job_id-named."""

    def test_sibling_of_pdf_path(self):
        pdf = tailored_pdf_path("ind-abc123")
        docx = tailored_docx_path("ind-abc123")
        # Same parent dir, different suffix
        assert docx.parent == pdf.parent
        assert pdf.suffix == ".pdf"
        assert docx.suffix == ".docx"

    def test_filename_does_not_contain_job_id(self):
        path = tailored_docx_path("ind-9934dbc8cae647b8")
        assert "ind-9934" not in path.name
        assert path.name.endswith("_Resume.docx") or path.name == "Resume.docx"

    def test_sanitizes_unsafe_chars_in_job_dir(self):
        path = tailored_docx_path("li/../etc/passwd")
        assert "/" not in path.parent.name
        assert "\\" not in path.parent.name


class TestRenderDocx:
    """The DOCX renderer produces a real .docx file with parseable
    content. We don't try to validate Word's full schema — we just
    confirm:
      - file is created
      - file is non-trivial size (would be ~0 bytes on render fail)
      - doc opens with python-docx and contains our text
    """

    def _resume(self):
        return TailoredResume(
            summary="Senior data analyst with 5 years of experience.",
            skills=["Python", "SQL", "Tableau"],
            experience=[{
                "title": "Senior Analyst",
                "company": "Acme Corp",
                "dates": "2021-Present",
                "bullets": [
                    "Built revenue dashboards used by exec team weekly",
                    "Led SQL training across 3 teams",
                ],
            }],
            education=[{
                "school": "State University",
                "degree": "BS Computer Science",
                "year": "2019",
            }],
        )

    def test_renders_real_docx_file(self, tmp_path):
        out = tmp_path / "test.docx"
        ok = asyncio.run(render_docx(
            self._resume(), out, name="Jane Doe", contact="jane@example.com",
        ))
        assert ok is True
        assert out.exists()
        assert out.stat().st_size > 1000  # real docx, not empty

    def test_docx_contains_expected_content(self, tmp_path):
        out = tmp_path / "test.docx"
        asyncio.run(render_docx(
            self._resume(), out, name="Jane Doe", contact="jane@example.com",
        ))
        from docx import Document
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Jane Doe" in all_text
        assert "jane@example.com" in all_text
        assert "Senior data analyst" in all_text
        # Skills joined as prose, not bullets, for parser-friendliness
        assert "Python, SQL, Tableau" in all_text
        # Experience block
        assert "Senior Analyst" in all_text
        assert "Acme Corp" in all_text
        assert "Built revenue dashboards" in all_text
        # Education block
        assert "State University" in all_text
        assert "BS Computer Science" in all_text

    def test_docx_uses_section_headings_uppercase(self, tmp_path):
        """Section headings should be uppercase ('SUMMARY', 'EXPERIENCE',
        etc.) to match the PDF template's visual style."""
        out = tmp_path / "test.docx"
        asyncio.run(render_docx(
            self._resume(), out, name="Jane", contact="",
        ))
        from docx import Document
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "SUMMARY" in all_text
        assert "SKILLS" in all_text
        assert "EXPERIENCE" in all_text
        assert "EDUCATION" in all_text

    def test_render_docx_handles_empty_sections(self, tmp_path):
        """A resume with no experience or education shouldn't crash —
        just emit empty sections."""
        out = tmp_path / "test.docx"
        empty_resume = TailoredResume(
            summary="Just starting out.",
            skills=["enthusiasm"],
            experience=[],
            education=[],
        )
        ok = asyncio.run(render_docx(
            empty_resume, out, name="Jane", contact="",
        ))
        assert ok is True
        assert out.exists()

    def test_render_docx_returns_false_on_python_docx_missing(
        self, tmp_path, monkeypatch
    ):
        """Defensive: if python-docx is missing (it's a dep but a
        future build could omit), we return False and let the caller
        fall back to PDF rather than crash."""
        import builtins
        original_import = builtins.__import__

        def fake_import(name, *args, **kwargs):
            if name == "docx" or name.startswith("docx."):
                raise ImportError("pretend python-docx isn't installed")
            return original_import(name, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", fake_import)
        out = tmp_path / "test.docx"
        ok = asyncio.run(render_docx(
            self._resume(), out, name="Jane", contact="",
        ))
        assert ok is False
        assert not out.exists()
