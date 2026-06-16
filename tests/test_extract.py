"""Résumé → fact-bank extraction (Direction 1, Phase A) — contract tests.

We fake the LLM (a stub returning a known JSON dict) so these stay deterministic and offline.
Text extraction is tested on real .txt/.docx files; the live qwen3 extraction quality is a
separate (manual / eval-harness) concern.

Covers:
  * extract_text_from_file: .txt and .docx round-trip; unsupported extension raises;
  * extract_factbank: a clean payload → the right FactBank; blank text → empty bank, NO LLM call;
    résumé-derived fields populate but work-auth/EEO stay default (never extracted);
  * _coerce_factbank_dict: drops stray keys, normalizes lists, filters empty/None entries,
    survives a noisy model response (still constructs a FactBank);
  * merge_extracted: takes résumé fields from the extraction, PRESERVES the user-entered
    work-auth / sponsorship / EEO / relocation from the existing bank.
"""

from __future__ import annotations

import asyncio

import pytest

from auto_applier.resume.extract import (
    _coerce_factbank_dict,
    extract_factbank,
    extract_text_from_file,
    merge_extracted,
)
from auto_applier.resume.factbank import Contact, FactBank, WorkEntry


# --------------------------------------------------------------- fakes

class _StubLLM:
    """Async completion stub — returns a fixed dict and counts calls."""

    def __init__(self, payload: dict):
        self._payload = payload
        self.calls = 0
        self.last_system = None
        self.last_think = "unset"
        self.last_num_predict = "unset"

    async def complete_json(
        self, prompt: str, *, system: str = "",
        think: bool | None = None, num_predict: int | None = None,
    ) -> dict:
        self.calls += 1
        self.last_system = system
        self.last_think = think
        self.last_num_predict = num_predict
        return self._payload


_GOOD_PAYLOAD = {
    "contact": {
        "name": "Jane Doe", "email": "jane@example.com", "phone": "555-0100",
        "location": "Dallas, TX", "links": {"LinkedIn": "https://lnkd.in/jane"},
    },
    "work_history": [
        {"company": "Acme", "title": "Senior Data Analyst", "start": "2020-01",
         "end": "Present", "bullets": ["Built the reporting platform", "Saved $40K/year"]},
    ],
    "education": [
        {"institution": "UT Austin", "degree": "B.S. Computer Science",
         "field_of_study": "CS", "start": "2012", "end": "2016"},
    ],
    "skills": ["SQL", "Python", "Power BI"],
    "certifications": ["AWS Solutions Architect"],
    "allowed_metrics": ["saved $40K/year", "190+ tables"],
}


# --------------------------------------------------------------- text extraction

def test_extract_text_from_txt(tmp_path):
    p = tmp_path / "resume.txt"
    p.write_text("Jane Doe\nSenior Data Analyst at Acme\n", encoding="utf-8")
    text = extract_text_from_file(p)
    assert "Jane Doe" in text and "Acme" in text


def test_extract_text_from_docx(tmp_path):
    import docx

    d = docx.Document()
    d.add_paragraph("Jane Doe")
    d.add_paragraph("Senior Data Analyst at Acme")
    p = tmp_path / "resume.docx"
    d.save(str(p))

    text = extract_text_from_file(p)
    assert "Jane Doe" in text and "Senior Data Analyst" in text


def test_extract_text_unsupported_extension_raises(tmp_path):
    p = tmp_path / "resume.doc"     # legacy binary .doc — unsupported
    p.write_text("not really a doc", encoding="utf-8")
    with pytest.raises(ValueError):
        extract_text_from_file(p)


# --------------------------------------------------------------- extract_factbank

def test_extract_factbank_maps_payload_to_bank():
    stub = _StubLLM(_GOOD_PAYLOAD)
    bank = asyncio.run(extract_factbank("some résumé text", stub))

    assert stub.calls == 1
    assert stub.last_think is False          # extraction disables qwen3 thinking (API param)
    assert stub.last_num_predict is not None  # ...and bounds the output (anti-degeneration)
    assert bank.contact.name == "Jane Doe"
    assert bank.contact.links == {"LinkedIn": "https://lnkd.in/jane"}
    assert len(bank.work_history) == 1
    w = bank.work_history[0]
    assert (w.company, w.title, w.end) == ("Acme", "Senior Data Analyst", "Present")
    assert w.bullets == ["Built the reporting platform", "Saved $40K/year"]
    assert bank.skills == ["SQL", "Python", "Power BI"]
    assert bank.allowed_metrics == ["saved $40K/year", "190+ tables"]
    # never extracted from a résumé — must stay at their no-silent-default values
    assert bank.work_authorization == ""
    assert bank.requires_sponsorship is None
    assert bank.eeo == {}


def test_extract_factbank_blank_text_skips_llm():
    stub = _StubLLM(_GOOD_PAYLOAD)
    bank = asyncio.run(extract_factbank("   \n  ", stub))
    assert stub.calls == 0                 # no LLM call on empty input
    assert bank.contact.name == ""
    assert bank.work_history == []


# --------------------------------------------------------------- coercion robustness

def test_coerce_drops_stray_keys_and_filters_junk():
    raw = {
        "contact": {"name": "J", "email": "e@x.com", "twitter": "@drop",
                    "links": {"GH": "u", "": "no-key", "X": ""}},
        "work_history": [
            {"company": "A", "title": "T", "bullets": ["keep", "", None]},
            {"foo": "bar"},                # no company/title → filtered out
        ],
        "education": [{"institution": "U", "degree": "BS"}, {}],
        "skills": ["a", "", None, "b"],
        "junk_top_level": "ignored",
        "allowed_metrics": ["m1"],
        "work_authorization": "should be ignored here",  # extraction must not carry this
    }
    d = _coerce_factbank_dict(raw)

    assert "twitter" not in d["contact"]
    assert d["contact"]["links"] == {"GH": "u"}        # empty key + empty value dropped
    assert len(d["work_history"]) == 1                  # {"foo":"bar"} filtered
    assert d["work_history"][0]["bullets"] == ["keep"]  # "" and None dropped
    assert len(d["education"]) == 1                      # {} filtered
    assert d["skills"] == ["a", "b"]
    assert "work_authorization" not in d                # not a résumé-derived field
    # and the coerced dict always builds a FactBank
    bank = FactBank.from_dict(d)
    assert bank.work_authorization == ""                # from_dict default, not the raw string


# --------------------------------------------------------------- merge

def test_merge_extracted_preserves_user_entered_fields():
    existing = FactBank(
        contact=Contact(name="Old Name"),
        skills=["oldskill"],
        work_authorization="US citizen",
        requires_sponsorship=False,
        eeo={"gender": "Male"},
        relocation={"willing": ["United States"]},
    )
    extracted = FactBank(
        contact=Contact(name="New Name", email="new@x.com"),
        work_history=[WorkEntry("Acme", "Analyst", "2020", "Present", ["did things"])],
        skills=["SQL", "Python"],
        allowed_metrics=["saved $40K"],
    )

    merged = merge_extracted(existing, extracted)

    # résumé-derived fields come from the extraction
    assert merged.contact.name == "New Name"
    assert merged.contact.email == "new@x.com"
    assert merged.skills == ["SQL", "Python"]
    assert merged.work_history[0].company == "Acme"
    assert merged.allowed_metrics == ["saved $40K"]
    # user-entered fields are preserved from the existing bank (never clobbered)
    assert merged.work_authorization == "US citizen"
    assert merged.requires_sponsorship is False
    assert merged.eeo == {"gender": "Male"}
    assert merged.relocation == {"willing": ["United States"]}
