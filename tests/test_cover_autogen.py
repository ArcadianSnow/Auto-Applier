"""Cover-letter autogen — BUILD 5 (``av3 cover --generate[-all]``).

Covers the offline core (no LLM wire, no browser): the .docx renderer, single-job
generate_one (happy / no-clobber / force / empty-JD / guard-fail / LLM-error), the batch
backfill (score floor, DECIDED-only, only-missing, limit), and the no-AI-tells voice
contract on the prompt. Same fake-the-client pattern as test_optimize_worker.py.
"""

from __future__ import annotations

import asyncio
import sqlite3

from docx import Document

from auto_applier.config.settings import Settings
from auto_applier.db.repositories import JobRepo, ScoreRepo
from auto_applier.domain.models import Job, JobScore
from auto_applier.domain.state import JobState
from auto_applier.llm.prompts import GENERATE_COVER_LETTER
from auto_applier.resume.cover_autogen import (
    ERROR,
    GENERATED,
    SKIPPED_DEGENERATE,
    SKIPPED_EXISTING,
    SKIPPED_GUARD,
    SKIPPED_INVALID,
    SKIPPED_NO_DESCRIPTION,
    _strip_ai_tells,
    backfill,
    generate_one,
    render_cover_letter_docx,
)
from auto_applier.resume.factbank import Contact, FactBank, WorkEntry
from auto_applier.resume.generate import (
    CoverLetterGenerator,
    existing_job_cover,
    job_cover_upload_path,
)

_JD = "We need a SQL Server DBA who is strong in Python and owns ETL pipelines."

# A letter that only claims bank-supported tech (SQL Server, Python) → guard PASS.
_HAPPY_BODY = (
    "I spent five years as a Database Administrator at Acme, working day to day in "
    "SQL Server and Python.\n\n"
    "My work refactoring the billing database lines up with what your team needs.\n\n"
    "I would be glad to talk it through."
)
# Claims Kubernetes/Terraform — neither in the bank → guard flags → not written.
_FABRICATED_BODY = (
    "I have deep Kubernetes and Terraform expertise running production clusters."
)
# The monotone "I did X. I did Y. I did Z." list the user rejected (every sentence opens with
# "I"); only bank-supported tech so the fab guard passes — the *voice* backstop is what flags it.
_MONOTONE_BODY = (
    "I worked in SQL Server and Python at Acme for five years.\n\n"
    "I refactored the billing database. I owned the SQL Server backups. "
    "I built ETL pipelines in Python. I tuned the database for the team.\n\n"
    "I would welcome the chance to talk."
)
# The same facts, varied openers (At/My/That lead most sentences) — the draft to prefer.
_VARIED_BODY = (
    "At Acme I spent five years refactoring the billing database in SQL Server.\n\n"
    "My Python ETL pipelines replaced a manual export, and the SQL Server backups I "
    "owned ran clean for three years. That work maps to what your team needs.\n\n"
    "I would welcome the chance to talk."
)
# qwen3 sometimes role-plays the recruiter instead of writing the letter (Databricks, 2026-06-15).
# Passes the fab guard (no fabricated tech) and isn't degenerate — the meta-response guard rejects it.
_META_BODY = (
    "Joseph, your resume is impressive, but we need to see more about your experience "
    "with distributed systems and query optimization. Could you please provide "
    "additional details on these areas?"
)


# --------------------------------------------------------------- fakes

class _CoverLLM:
    """Minimal CompletionClient: returns a fixed cover body (or raises). Counts calls so
    the no-clobber / empty-JD short-circuits can be proven to skip the LLM."""

    def __init__(self, body: str = _HAPPY_BODY, *, raise_exc: Exception | None = None):
        self._body = body
        self._raise = raise_exc
        self.calls = 0

    async def complete_json(self, prompt: str, *, system: str = "") -> dict:
        self.calls += 1
        if self._raise is not None:
            raise self._raise
        return {"body": self._body}


class _SeqCoverLLM:
    """Returns a scripted sequence of bodies (last one repeats) so the third-person
    regenerate path can be driven deterministically."""

    def __init__(self, bodies: list[str]):
        self._bodies = list(bodies)
        self.calls = 0

    async def complete_json(self, prompt: str, *, system: str = "") -> dict:
        body = self._bodies[min(self.calls, len(self._bodies) - 1)]
        self.calls += 1
        return {"body": body}


class _FlakyCoverLLM:
    """Raises on the first call (a timeout), succeeds on the second — drives the no_think
    fail-safe retry. Records each prompt so the test can prove '/no_think' is appended only
    on the retry."""

    def __init__(self, body: str = _HAPPY_BODY, exc: Exception | None = None):
        self._body = body
        self._exc = exc or RuntimeError("ReadTimeout")
        self.calls = 0
        self.prompts: list[str] = []

    async def complete_json(self, prompt: str, *, system: str = "") -> dict:
        self.calls += 1
        self.prompts.append(prompt)
        if self.calls == 1:
            raise self._exc
        return {"body": self._body}


def _bank() -> FactBank:
    return FactBank(
        contact=Contact(
            name="Joseph Lira", email="jl@example.com",
            phone="+1 555 0100", location="Dallas, TX",
        ),
        skills=["sql", "sql server", "python", "etl"],
        work_history=[
            WorkEntry(
                company="Acme", title="Database Administrator",
                start="2019", end="2024",
                bullets=["Refactored the billing database", "Owned SQL Server backups"],
            )
        ],
        allowed_metrics=[],
    )


def _job(jid: str = "j1", *, description: str = _JD,
         company: str = "BetaCo", title: str = "Data Platform Engineer") -> Job:
    return Job(
        id=jid, source="greenhouse", source_job_id=jid, title=title,
        company=company, description=description, state=JobState.DECIDED,
    )


def _gen_one(settings: Settings, job: Job, llm: _CoverLLM, *, name: str = "Joseph Lira",
             force: bool = False):
    return asyncio.run(generate_one(
        settings, job, bank=_bank(), generator=CoverLetterGenerator(llm),
        name=name, force=force,
    ))


def _seed_scored(conn: sqlite3.Connection, *, sid: str, total: float,
                 state: JobState = JobState.DECIDED, description: str = _JD,
                 company: str = "BetaCo", title: str = "Data Platform Engineer") -> str:
    repo = JobRepo(conn)
    job = Job(source="greenhouse", source_job_id=sid, title=title,
              company=company, description=description)
    repo.add(job)
    repo.set_state(job.id, JobState.DESCRIBED)
    repo.set_state(job.id, JobState.SCORED)
    repo.set_state(job.id, JobState.DECIDED)
    if state is JobState.REVIEW:
        repo.set_state(job.id, JobState.REVIEW)
    elif state is JobState.SKIPPED:
        repo.set_state(job.id, JobState.SKIPPED)
    ScoreRepo(conn).upsert(JobScore(job_id=job.id, total=total))
    return job.id


# --------------------------------------------------------------- .docx render

def test_render_docx_is_a_complete_letter(settings, tmp_path):
    out = tmp_path / "letter.docx"
    render_cover_letter_docx(_HAPPY_BODY, _bank().contact, out)
    assert out.exists()

    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    joined = "\n".join(texts)
    assert "Dear Hiring Manager," in texts
    assert "Sincerely," in texts
    assert texts.count("Joseph Lira") == 2  # header + signature
    assert "refactoring the billing database" in joined
    assert "jl@example.com | +1 555 0100 | Dallas, TX" in texts
    # author metadata is the applicant's name, not the python-docx default tell
    assert doc.core_properties.author == "Joseph Lira"


def test_render_docx_no_contact_still_renders(settings, tmp_path):
    out = tmp_path / "bare.docx"
    render_cover_letter_docx(_HAPPY_BODY, Contact(), out)
    texts = [p.text for p in Document(str(out)).paragraphs]
    assert "Dear Hiring Manager," in texts
    assert "Sincerely," in texts


# --------------------------------------------------------------- generate_one

def test_generate_one_happy_writes_docx(settings):
    llm = _CoverLLM()
    res = _gen_one(settings, _job("happy"), llm)

    assert res.status == GENERATED and res.ok
    assert llm.calls == 1
    dest = job_cover_upload_path(settings, "happy", ".docx", "Joseph Lira")
    assert dest.exists() and res.path == str(dest)
    assert dest.name == "Joseph Lira Cover Letter.docx"
    # the apply path / av3 cover would now find it
    assert existing_job_cover(settings, "happy") == dest


def test_generate_one_no_clobber_skips_and_skips_llm(settings):
    # a hand-authored letter is already assigned
    manual = job_cover_upload_path(settings, "manual", ".docx", "")
    manual.parent.mkdir(parents=True, exist_ok=True)
    manual.write_text("HAND WRITTEN", encoding="utf-8")

    llm = _CoverLLM()
    res = _gen_one(settings, _job("manual"), llm)

    assert res.status == SKIPPED_EXISTING
    assert llm.calls == 0  # never spent an LLM call
    assert manual.read_text(encoding="utf-8") == "HAND WRITTEN"  # untouched


def test_generate_one_force_overwrites_single_file(settings):
    manual = job_cover_upload_path(settings, "f", ".docx", "")
    manual.parent.mkdir(parents=True, exist_ok=True)
    manual.write_text("OLD", encoding="utf-8")

    res = _gen_one(settings, _job("f"), _CoverLLM(), force=True)
    assert res.status == GENERATED

    folder = manual.parent
    covers = sorted(p.name for p in folder.glob("*Cover Letter.*"))
    assert covers == ["Joseph Lira Cover Letter.docx"]  # exactly one, the new one
    # and it's a real docx now, not the OLD text
    assert "Dear Hiring Manager," in [p.text for p in Document(str(folder / covers[0])).paragraphs]


def test_generate_one_empty_description_skips(settings):
    llm = _CoverLLM()
    res = _gen_one(settings, _job("nodesc", description="   "), llm)
    assert res.status == SKIPPED_NO_DESCRIPTION
    assert llm.calls == 0


def test_generate_one_guard_fail_writes_nothing(settings):
    res = _gen_one(settings, _job("guard"), _CoverLLM(body=_FABRICATED_BODY))
    assert res.status == SKIPPED_GUARD
    assert "kubernetes" in res.detail.lower() or "terraform" in res.detail.lower()
    assert existing_job_cover(settings, "guard") is None  # nothing shipped


def test_generate_one_llm_error_is_isolated(settings):
    res = _gen_one(settings, _job("err"), _CoverLLM(raise_exc=RuntimeError("ollama down")))
    assert res.status == ERROR
    assert "ollama down" in res.detail
    assert existing_job_cover(settings, "err") is None


# --------------------------------------------------------------- dash strip (the #1 AI tell)

def test_strip_ai_tells_replaces_dashes():
    assert _strip_ai_tells("I did A—then B.") == "I did A, then B."
    assert _strip_ai_tells("word–word") == "word, word"
    assert "—" not in _strip_ai_tells("a — b — c")
    # paragraph breaks survive the strip (newline-safe)
    assert "\n\n" in _strip_ai_tells("Para one—done.\n\nPara two.")


def test_generate_one_strips_em_dash_even_if_llm_emits_one(settings):
    body = "I ran SQL Server—and Python—at Acme.\n\nGlad to talk."
    res = _gen_one(settings, _job("dash"), _CoverLLM(body=body))
    assert res.status == GENERATED
    txt = "\n".join(p.text for p in Document(res.path).paragraphs)
    assert "—" not in txt and "–" not in txt


# --------------------------------------------------------------- backfill

def test_backfill_only_strong_decided_jobs(settings, conn):
    strong = _seed_scored(conn, sid="strong", total=9.0)
    mid = _seed_scored(conn, sid="mid", total=8.2)
    weak = _seed_scored(conn, sid="weak", total=6.0)  # below floor — never touched

    results = asyncio.run(backfill(
        settings, conn, llm=_CoverLLM(), bank=_bank(), min_score=8.0, name="Joseph Lira",
    ))

    gen_ids = {r.job_id for r in results if r.status == GENERATED}
    assert gen_ids == {strong, mid}
    assert existing_job_cover(settings, strong) is not None
    assert existing_job_cover(settings, mid) is not None
    assert existing_job_cover(settings, weak) is None
    assert all(r.job_id != weak for r in results)


def test_backfill_skips_existing_and_respects_state(settings, conn):
    has_letter = _seed_scored(conn, sid="has", total=9.5)
    review = _seed_scored(conn, sid="rev", total=9.0, state=JobState.REVIEW)  # not DECIDED
    fresh = _seed_scored(conn, sid="fresh", total=8.5)

    # pre-assign a manual letter to `has`
    p = job_cover_upload_path(settings, has_letter, ".docx", "")
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text("MANUAL", encoding="utf-8")

    results = asyncio.run(backfill(
        settings, conn, llm=_CoverLLM(), bank=_bank(), min_score=8.0, name="Joseph Lira",
    ))
    by_id = {r.job_id: r.status for r in results}

    assert by_id[has_letter] == SKIPPED_EXISTING
    assert p.read_text(encoding="utf-8") == "MANUAL"  # never clobbered
    assert by_id[fresh] == GENERATED
    assert review not in by_id  # REVIEW state filtered out (DECIDED-only)


def test_backfill_limit_caps_new_letters(settings, conn):
    _seed_scored(conn, sid="a", total=9.9)
    _seed_scored(conn, sid="b", total=9.8)
    _seed_scored(conn, sid="c", total=9.7)

    results = asyncio.run(backfill(
        settings, conn, llm=_CoverLLM(), bank=_bank(), min_score=8.0, limit=1,
    ))
    assert len([r for r in results if r.status == GENERATED]) == 1


def test_backfill_empty_when_nothing_qualifies(settings, conn):
    _seed_scored(conn, sid="low", total=5.0)
    results = asyncio.run(backfill(settings, conn, llm=_CoverLLM(), bank=_bank(), min_score=8.0))
    assert results == []


# --------------------------------------------------------------- voice contract

def test_cover_prompt_enforces_no_ai_tells_voice():
    assert GENERATE_COVER_LETTER.version == "gen-cover-v4"
    sys = GENERATE_COVER_LETTER.system.lower()
    assert "em-dash" in sys
    assert "excited" in sys and "thrilled" in sys      # categorical "excited" family ban
    assert "passionate" in sys                          # a banned buzzword is named
    assert "rule of three" in sys
    # anti-overclaim: the prompt must forbid inventing soft experience the bank lacks
    assert "overclaim" in sys or "do not claim" in sys
    # v4: kill the robotic "I did X. I did Y." list — at most one "I" opener, with a
    # concrete BAD/GOOD rhythm example (qwen3 follows examples better than abstract rules)
    assert "do not write a list" in sys
    assert "begin with the word 'i'" in sys
    assert "bad" in sys and "good" in sys
    # v4: don't parrot the JD's marketing adjectives back
    assert "parrot" in sys and "scalable" in sys
    # v4: exactly three paragraphs, at most two accomplishments (a letter, not a résumé)
    assert "exactly three short paragraphs" in sys
    assert "at most two" in sys


def test_trim_jd_for_cover_caps_big_jd_at_word_boundary():
    from auto_applier.resume.generate import _COVER_JD_MAX_CHARS, _trim_jd_for_cover

    short = "We need a SQL Server DBA."
    assert _trim_jd_for_cover(short) == short        # short JD passes through untouched

    big = ("word " * 4000).strip()                   # ~20k chars, well over the cap
    trimmed = _trim_jd_for_cover(big)
    assert trimmed.endswith("[...]")                 # marked as truncated
    assert len(trimmed) <= _COVER_JD_MAX_CHARS + 8   # body capped (plus the short marker)
    assert "wordword" not in trimmed                 # cut at a whitespace boundary, no split word


def test_opens_in_third_person_detects_name_and_pronoun():
    from auto_applier.resume.cover_autogen import _opens_in_third_person

    assert _opens_in_third_person("Joseph Lira has built data systems.", "Joseph Lira")
    assert _opens_in_third_person("Joseph has built data systems.", "Joseph Lira")
    assert _opens_in_third_person("He has built data systems.", "Joseph Lira")
    assert _opens_in_third_person("His work spans SQL.", "Joseph Lira")
    assert not _opens_in_third_person("I built data systems.", "Joseph Lira")
    assert not _opens_in_third_person("At Acme, I built data systems.", "Joseph Lira")
    assert not _opens_in_third_person("", "Joseph Lira")


def test_generate_one_regenerates_past_a_third_person_opening(settings):
    # First draft opens in the third person (a defect); generate_one must regenerate once
    # and accept the first-person second draft. Both bodies claim only bank-supported tech.
    third = ("Joseph Lira has worked in SQL Server and Python at Acme.\n\n"
             "He refactored the billing database.\n\nGlad to talk.")
    llm = _SeqCoverLLM([third, _HAPPY_BODY])
    res = asyncio.run(generate_one(
        settings, _job(), bank=_bank(), generator=CoverLetterGenerator(llm),
        name="Joseph Lira",
    ))
    assert res.status == GENERATED
    assert llm.calls == 2  # one retry consumed the third-person draft


def test_generate_one_retries_without_thinking_after_a_timeout(settings):
    # First attempt errors (simulating a JD whose reasoning blew past the read timeout); the
    # fail-safe second attempt drops qwen3 thinking via /no_think and succeeds.
    llm = _FlakyCoverLLM()
    res = asyncio.run(generate_one(
        settings, _job("slow"), bank=_bank(), generator=CoverLetterGenerator(llm),
        name="Joseph Lira",
    ))
    assert res.status == GENERATED
    assert llm.calls == 2
    assert "/no_think" not in llm.prompts[0]   # first attempt keeps thinking on
    assert "/no_think" in llm.prompts[1]        # the retry drops it


def test_generate_one_rejects_degenerate_output(settings):
    # qwen3 can loop and repeat a sentence hundreds of times (Mistral JDs, 2026-06-15). The
    # claims are all bank-supported so the fab guard would pass it — the degeneracy gate must
    # reject it first so a 10K-char monstrosity never ships as a "ready" letter.
    runaway = "I built a data pipeline in SQL Server and Python. " * 200
    llm = _CoverLLM(body=runaway)
    res = _gen_one(settings, _job("loop"), llm)
    assert res.status == SKIPPED_DEGENERATE
    assert existing_job_cover(settings, "loop") is None  # nothing shipped
    assert llm.calls == 2  # regenerated once, still degenerate → rejected


def test_ensure_paragraphs_splits_a_dense_block_into_three():
    from auto_applier.resume.cover_autogen import _ensure_paragraphs

    dense = "I built A. I designed B. I engineered C. I would welcome a talk."
    parts = _ensure_paragraphs(dense).split("\n\n")
    assert len(parts) == 3
    assert parts[0] == "I built A."            # hook = first sentence
    assert parts[2] == "I would welcome a talk."  # close = last sentence
    assert parts[1] == "I designed B. I engineered C."  # middle = the rest

    # a body the model already paragraphed is left exactly as-is
    para = "One.\n\nTwo two.\n\nThree."
    assert _ensure_paragraphs(para) == para
    # too short to regroup meaningfully → untouched
    short = "I built A. I would welcome a talk."
    assert _ensure_paragraphs(short) == short


def test_excessive_i_openings_flags_the_i_list_voice():
    from auto_applier.resume.cover_autogen import _excessive_i_openings, _i_opening_ratio

    monotone = "I built A. I designed B. I implemented C. I shipped D. I tested E."
    assert _excessive_i_openings(monotone)
    assert _i_opening_ratio(monotone) == 1.0
    # "In/It/If" openers are NOT "I" openers — they read as varied, good
    varied = ("At Acme I led the work. In practice it shipped on time. "
              "My team owned the rollout. I would welcome a talk.")
    assert not _excessive_i_openings(varied)
    assert _i_opening_ratio(varied) == 0.25  # only the last sentence opens with "I"
    # three "I" openers in a row trips it even in a short letter
    run3 = "My role grew fast. I built A. I designed B. I shipped C."
    assert _excessive_i_openings(run3)
    # one or two "I" openers in a normal letter is fine
    assert not _excessive_i_openings(
        "At Acme, I built data systems. The work shipped. I would welcome a talk."
    )
    assert not _excessive_i_openings("")  # empty isn't "monotone"


def test_generate_one_prefers_the_less_monotone_draft(settings):
    # First draft is the monotone "I did X. I did Y. I did Z." list the user rejects; the
    # generator must regenerate (no_think) and keep the varied second draft. Both bodies claim
    # only bank-supported tech, so the fab guard passes either — the voice backstop decides.
    llm = _SeqCoverLLM([_MONOTONE_BODY, _VARIED_BODY])
    res = asyncio.run(generate_one(
        settings, _job(), bank=_bank(), generator=CoverLetterGenerator(llm),
        name="Joseph Lira",
    ))
    assert res.status == GENERATED
    assert llm.calls == 2  # the monotone first draft triggered the retry
    txt = "\n".join(p.text for p in Document(res.path).paragraphs)
    assert "At Acme I spent five years" in txt  # the varied draft is the one that shipped


def test_generate_one_keeps_monotone_draft_if_retry_is_no_better(settings):
    # If both drafts are monotone (qwen3 just won't vary on this JD), still ship a letter —
    # a monotone-but-honest letter beats none. The voice backstop never *skips* like degeneracy.
    llm = _SeqCoverLLM([_MONOTONE_BODY, _MONOTONE_BODY])
    res = asyncio.run(generate_one(
        settings, _job("mono"), bank=_bank(), generator=CoverLetterGenerator(llm),
        name="Joseph Lira",
    ))
    assert res.status == GENERATED
    assert existing_job_cover(settings, "mono") is not None


def test_is_meta_response_flags_recruiter_reply_not_normal_letter():
    from auto_applier.resume.cover_autogen import _is_meta_response

    assert _is_meta_response(_META_BODY, "Joseph Lira")                       # the live hallucination
    assert _is_meta_response("Joseph Lira: thanks for applying.", "Joseph Lira")  # name as greeting
    assert _is_meta_response("Tell us more about your experience with Spark.", "Joseph Lira")
    # a normal letter that says "your team / your platform" (refers to the COMPANY) is fine
    assert not _is_meta_response(
        "At Acme I built ETL pipelines. This fits your team's focus on data quality. "
        "I'd welcome a talk.", "Joseph Lira")
    assert not _is_meta_response(_HAPPY_BODY, "Joseph Lira")
    assert not _is_meta_response("", "Joseph Lira")


def test_generate_one_skips_a_meta_response(settings):
    # The "Joseph, your resume is impressive..." recruiter-reply hallucination passes the fab
    # guard and isn't degenerate, so a dedicated guard must reject it — never ship a non-letter.
    llm = _CoverLLM(body=_META_BODY)
    res = _gen_one(settings, _job("meta"), llm)
    assert res.status == SKIPPED_INVALID
    assert existing_job_cover(settings, "meta") is None  # nothing shipped
    assert llm.calls == 2  # regenerated once, still meta → rejected


def test_generate_one_regenerates_past_a_meta_response(settings):
    # A meta first draft must trigger the retry and be replaced by the real second draft.
    llm = _SeqCoverLLM([_META_BODY, _HAPPY_BODY])
    res = asyncio.run(generate_one(
        settings, _job("metaretry"), bank=_bank(), generator=CoverLetterGenerator(llm),
        name="Joseph Lira",
    ))
    assert res.status == GENERATED
    assert llm.calls == 2


def test_is_degenerate_flags_runaway_and_repetition():
    from auto_applier.resume.cover_autogen import _is_degenerate

    assert not _is_degenerate(_HAPPY_BODY)                 # a real letter is fine
    assert _is_degenerate("I built a data pipeline. " * 200)  # absurd length (qwen3 loop)
    # same substantial sentence repeated → looping, even under the length cap
    loop = "I designed scalable upsert frameworks across many tables. " * 4
    assert _is_degenerate(loop)
    # a wall of 13+ DISTINCT substantial sentences is a list-dump, not a letter
    dump = " ".join(f"I built a distinct data system numbered {i} for the team." for i in range(15))
    assert _is_degenerate(dump)
    assert not _is_degenerate("")                          # empty isn't "degenerate"
