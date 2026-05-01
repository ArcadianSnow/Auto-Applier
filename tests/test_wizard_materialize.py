"""Integration tests for the wizard's eager-materialize path.

These exist because the materialize functions in
``auto_applier/gui/steps/resumes.py`` and
``auto_applier/gui/wizard.py`` import from auto_applier.resume.parser
at function-body time — so a typo in the imported name (e.g.
``parse_resume`` when the function is actually ``extract_text``)
manifests as a silent ``ImportError`` swallowed by tk's main loop.
The user's only signal is "the button does nothing".

The fix landed: rename to ``extract_text``. These tests guard
against the regression by importing the module at collection time
(so any broken import fails the test discovery) AND by exercising
the materialize call against a temp directory tree.
"""
from __future__ import annotations

import json
from pathlib import Path
from unittest.mock import MagicMock

import pytest


# Module-import guard: if the gui modules can't be loaded (because
# of a broken import like the parse_resume bug), pytest collection
# fails noisily instead of silently. That's the canary.
def test_resumes_step_module_imports_clean():
    from auto_applier.gui.steps import resumes  # noqa: F401


def test_wizard_module_imports_clean():
    from auto_applier.gui import wizard  # noqa: F401


@pytest.fixture
def tmp_data_dirs(tmp_path, monkeypatch):
    """Redirect RESUMES_DIR and PROFILES_DIR into a temp tree so the
    materialize call has a clean target without polluting real data/."""
    resumes = tmp_path / "data" / "resumes"
    profiles = tmp_path / "data" / "profiles"
    import auto_applier.config as cfg
    monkeypatch.setattr(cfg, "RESUMES_DIR", resumes)
    monkeypatch.setattr(cfg, "PROFILES_DIR", profiles)
    return resumes, profiles


@pytest.fixture
def fake_docx(tmp_path):
    """Create a minimal valid .docx so extract_text doesn't fail.

    A real DOCX is a zip with specific internal structure. Easiest
    way to get one in a test: use python-docx itself.
    """
    from docx import Document
    path = tmp_path / "fake_resume.docx"
    doc = Document()
    doc.add_paragraph("Fake Person")
    doc.add_paragraph("Senior Test Subject at Acme")
    doc.save(str(path))
    return path


def test_materialize_resume_copies_file_and_writes_profile(
    tmp_data_dirs, fake_docx, monkeypatch
):
    """The wizard's eager materialize must copy the source resume
    into RESUMES_DIR and write a profile JSON into PROFILES_DIR."""
    from auto_applier.gui.steps.resumes import ResumesStep

    resumes_dir, profiles_dir = tmp_data_dirs

    # Construct a ResumesStep without actually building the GUI —
    # we only need _materialize_resume's behaviour, not its widgets.
    step = ResumesStep.__new__(ResumesStep)
    step.wizard = MagicMock()

    ok = step._materialize_resume(str(fake_docx), "MyLabel", verbose=False)

    assert ok is True
    # File copied with the label as basename.
    assert (resumes_dir / "MyLabel.docx").exists()
    # Profile JSON written with extracted text + the label.
    profile_path = profiles_dir / "MyLabel.json"
    assert profile_path.exists()
    profile = json.loads(profile_path.read_text(encoding="utf-8"))
    assert profile["label"] == "MyLabel"
    assert profile["source_file"] == "MyLabel.docx"
    assert "Fake Person" in profile["raw_text"]
    assert profile["skills"] == []
    assert profile["confirmed_skills"] == []


def test_materialize_resume_handles_missing_source(
    tmp_data_dirs, tmp_path
):
    from auto_applier.gui.steps.resumes import ResumesStep

    step = ResumesStep.__new__(ResumesStep)
    step.wizard = MagicMock()
    bogus = tmp_path / "does_not_exist.docx"
    ok = step._materialize_resume(str(bogus), "Ghost", verbose=False)
    assert ok is False
    resumes_dir, profiles_dir = tmp_data_dirs
    assert not (resumes_dir / "Ghost.docx").exists()
    assert not (profiles_dir / "Ghost.json").exists()


def test_materialize_resume_sanitizes_label_with_unsafe_chars(
    tmp_data_dirs, fake_docx
):
    """Labels containing path-unsafe chars (/, \\, :, *, ?, etc.)
    used to silently fail with OSError 22 on Windows. The sanitizer
    rewrites them to underscores so the copy works."""
    from auto_applier.gui.steps.resumes import ResumesStep

    resumes_dir, _ = tmp_data_dirs
    step = ResumesStep.__new__(ResumesStep)
    step.wizard = MagicMock()
    ok = step._materialize_resume(
        str(fake_docx), "Data: Analyst*?", verbose=False,
    )
    assert ok is True
    # All offending chars become underscore.
    files = list(resumes_dir.glob("*.docx"))
    assert len(files) == 1
    assert "/" not in files[0].name
    assert ":" not in files[0].name
    assert "*" not in files[0].name
    assert "?" not in files[0].name


def test_materialize_pending_resumes_handles_full_list(
    tmp_data_dirs, fake_docx, tmp_path
):
    """wizard._materialize_pending_resumes iterates resume_list and
    materializes each — covers the case where the user added resumes
    on an old build (pre-eager-copy) and now the wizard saves on Next."""
    from auto_applier.gui.wizard import WizardApp

    resumes_dir, profiles_dir = tmp_data_dirs

    # Build a wizard-like stand-in with just the bits we need.
    wiz = WizardApp.__new__(WizardApp)
    wiz.resume_list = [
        ("FirstResume", str(fake_docx)),
        ("SecondResume", str(fake_docx)),
    ]
    wiz._materialize_pending_resumes()

    assert (resumes_dir / "FirstResume.docx").exists()
    assert (resumes_dir / "SecondResume.docx").exists()
    assert (profiles_dir / "FirstResume.json").exists()
    assert (profiles_dir / "SecondResume.json").exists()
