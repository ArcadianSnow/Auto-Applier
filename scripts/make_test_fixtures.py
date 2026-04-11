"""Generate a complete fake-persona test fixture set for dry-run testing.

Creates:
- data/resumes/testpilot.docx         — a parseable DOCX resume
- data/user_config.json               — personal info the wizard would populate
- data/answers.json                   — pre-canned answers for common form fields
- data/archetypes.json                — example archetypes for routing (opt-in)
- data/research/sample_source.txt     — source material to try `cli research`

Everything lands under data/ which is gitignored, so nothing leaks into
commits. Safe to re-run — overwrites existing test files.

Run: python scripts/make_test_fixtures.py
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
DATA = ROOT / "data"


# ---------------------------------------------------------------------------
# The fake persona
# ---------------------------------------------------------------------------

PERSONA = {
    "name": "Jordan Testpilot",
    "first_name": "Jordan",
    "last_name": "Testpilot",
    "email": "jordan.testpilot@example.com",
    "phone": "+1 555 0100",
    "city": "Seattle, WA",
    "linkedin_url": "https://www.linkedin.com/in/jordan-testpilot-fake",
    "github_url": "https://github.com/jordan-testpilot-fake",
    "portfolio_url": "https://jordan-testpilot.example.com",
    "work_auth": "US Citizen",
    "requires_sponsorship": "No",
    "willing_to_relocate": "Yes",
    "desired_salary": "120000",
    "years_experience": "6",
}


# ---------------------------------------------------------------------------
# Resume DOCX
# ---------------------------------------------------------------------------

def write_resume() -> Path:
    doc = Document()

    doc.add_heading(PERSONA["name"], level=0)
    doc.add_paragraph(
        f"{PERSONA['email']} | {PERSONA['phone']} | {PERSONA['city']}\n"
        f"{PERSONA['linkedin_url']} | {PERSONA['github_url']}"
    )

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "Data analyst with 6 years of experience building dashboards, "
        "automating reporting pipelines, and translating ambiguous business "
        "questions into SQL. Comfortable owning a project end-to-end from "
        "stakeholder discovery through production deployment."
    )

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(
        "Python · SQL (PostgreSQL, BigQuery, Snowflake) · pandas · dbt · "
        "Airflow · Tableau · Looker · Git · Docker · AWS (S3, Athena, "
        "Lambda) · Excel · A/B testing · statistical analysis"
    )

    doc.add_heading("Experience", level=1)

    doc.add_heading("Senior Data Analyst — Northwind Logistics", level=2)
    doc.add_paragraph("Seattle, WA · Mar 2022 – Present")
    for bullet in [
        "Owned the analytics stack for a $40M freight brokerage product, "
        "including 30+ dbt models feeding Looker dashboards consumed by "
        "operations and sales leadership.",
        "Built a driver-retention forecasting pipeline in BigQuery + "
        "Python that reduced churn-related revenue loss by 14% over the "
        "following two quarters.",
        "Partnered with a cross-functional team of 4 engineers and 2 PMs "
        "to instrument a new pricing engine, landing on a shared metrics "
        "definition that cut reporting disputes by half.",
        "Mentored two junior analysts on SQL style, dbt testing, and "
        "stakeholder communication — both promoted within a year.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Data Analyst — Brightwire Energy", level=2)
    doc.add_paragraph("Portland, OR · Jun 2019 – Feb 2022")
    for bullet in [
        "Migrated weekly reporting from Excel to Tableau, cutting manual "
        "prep time from 8 hours to 20 minutes and enabling self-serve "
        "exploration for three downstream teams.",
        "Designed the company's first A/B testing framework for a "
        "residential solar landing page, which informed a redesign that "
        "improved lead-to-quote conversion by 22%.",
        "Automated monthly board-deck metrics with a Python + Snowflake "
        "pipeline orchestrated in Airflow, eliminating a recurring "
        "5-hour pre-meeting rush.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Junior Analyst — Citywide Credit Union", level=2)
    doc.add_paragraph("Tacoma, WA · Aug 2017 – May 2019")
    for bullet in [
        "Produced member-growth and loan-portfolio reports using SQL "
        "Server and Excel for a credit union serving 80,000 members.",
        "Built an early-warning flag for delinquent-loan risk that "
        "surfaced 30% more at-risk accounts than the manual process "
        "it replaced.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Education", level=1)
    doc.add_paragraph(
        "B.S. Statistics — University of Washington, 2017"
    )

    path = DATA / "resumes" / "testpilot.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# user_config.json
# ---------------------------------------------------------------------------

def write_user_config() -> Path:
    config = {
        "personal_info": {
            "name": PERSONA["name"],
            "first_name": PERSONA["first_name"],
            "last_name": PERSONA["last_name"],
            "email": PERSONA["email"],
            "phone": PERSONA["phone"],
            "city": PERSONA["city"],
            "linkedin_url": PERSONA["linkedin_url"],
            "github_url": PERSONA["github_url"],
            "portfolio_url": PERSONA["portfolio_url"],
            "work_auth": PERSONA["work_auth"],
            "requires_sponsorship": PERSONA["requires_sponsorship"],
            "willing_to_relocate": PERSONA["willing_to_relocate"],
            "desired_salary": PERSONA["desired_salary"],
            "years_experience": PERSONA["years_experience"],
        },
        "search_keywords": [
            "data analyst",
            "analytics engineer",
            "business intelligence analyst",
        ],
        "search_location": "Seattle, WA",
        "enabled_platforms": ["linkedin", "indeed"],
        "dry_run": True,
        "max_applications_per_day": 3,
        "auto_apply_min": 7,
        "cli_auto_apply_min": 7,
        "review_min": 4,
        "ollama_model": "gemma4:e4b",
    }
    path = DATA / "user_config.json"
    path.write_text(json.dumps(config, indent=2), encoding="utf-8")
    return path


# ---------------------------------------------------------------------------
# answers.json
# ---------------------------------------------------------------------------

def write_answers() -> Path:
    answers = {
        "questions": [
            {
                "question": "Are you legally authorized to work in the United States?",
                "answer": "Yes",
                "aliases": [
                    "work authorization",
                    "authorized to work",
                    "legally authorized",
                ],
            },
            {
                "question": "Will you now or in the future require sponsorship for employment visa status?",
                "answer": "No",
                "aliases": ["require sponsorship", "visa sponsorship"],
            },
            {
                "question": "How many years of experience do you have with SQL?",
                "answer": "6",
                "aliases": ["years of sql", "sql experience"],
            },
            {
                "question": "How many years of experience do you have with Python?",
                "answer": "5",
                "aliases": ["years of python", "python experience"],
            },
            {
                "question": "Are you willing to relocate?",
                "answer": "Yes",
                "aliases": ["relocate", "relocation"],
            },
            {
                "question": "What are your salary expectations?",
                "answer": "120000",
                "aliases": ["salary expectation", "desired salary", "compensation"],
            },
            {
                "question": "Do you have a Bachelor's degree or higher?",
                "answer": "Yes",
                "aliases": ["bachelor", "degree"],
            },
        ]
    }
    path = DATA / "answers.json"
    path.write_text(json.dumps(answers, indent=2), encoding="utf-8")
    return path


# ---------------------------------------------------------------------------
# archetypes.json (opt-in routing)
# ---------------------------------------------------------------------------

def write_archetypes() -> Path:
    archetypes = {
        "archetypes": [
            {
                "name": "data_analyst",
                "description": (
                    "SQL, dashboards, business reporting, stakeholder "
                    "partnership, basic statistics"
                ),
                "keywords": ["sql", "tableau", "looker", "dashboards", "analyst"],
            },
            {
                "name": "analytics_engineer",
                "description": (
                    "dbt, data modeling, pipeline ownership, testing, "
                    "git-based analytics workflows"
                ),
                "keywords": ["dbt", "data modeling", "analytics engineer"],
            },
            {
                "name": "data_engineer",
                "description": (
                    "Airflow, Spark, streaming, data infrastructure, "
                    "production pipelines"
                ),
                "keywords": ["airflow", "spark", "kafka", "etl", "pipeline"],
            },
        ]
    }
    path = DATA / "archetypes.json"
    path.write_text(json.dumps(archetypes, indent=2), encoding="utf-8")
    return path


# ---------------------------------------------------------------------------
# Sample research source material
# ---------------------------------------------------------------------------

def write_research_sample() -> Path:
    material = (
        "About Northwind Logistics\n"
        "Northwind is a Seattle-based freight brokerage platform founded in "
        "2015. We connect independent truck drivers with shippers through a "
        "mobile-first marketplace that handles load matching, invoicing, and "
        "settlement. We've grown to 40 employees and moved $40M in freight "
        "last year.\n\n"
        "Our Engineering Culture\n"
        "We ship small, we ship often, and we trust engineers to own their "
        "work end-to-end. No on-call for the analytics team — that's handled "
        "by platform infra. Every feature starts with a written one-pager. "
        "Our stack is Python, PostgreSQL, dbt, Airflow, and React on the "
        "front end. We're remote-first with an optional office in Pioneer "
        "Square.\n\n"
        "Open Roles\n"
        "- Senior Analytics Engineer (remote, $140-170k)\n"
        "- Staff Data Scientist (remote, $170-210k)\n"
        "- Principal Platform Engineer (hybrid, $190-230k)\n\n"
        "Recent news: Raised a $25M Series B in early 2026 led by Benchmark "
        "with participation from Sequoia. Plans to expand into Canadian "
        "markets by Q3."
    )
    path = DATA / "research" / "sample_source.txt"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(material, encoding="utf-8")
    return path


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    print("Generating test fixtures under data/ (gitignored)...")
    resume = write_resume()
    print(f"  wrote {resume}")
    cfg = write_user_config()
    print(f"  wrote {cfg}")
    ans = write_answers()
    print(f"  wrote {ans}")
    arch = write_archetypes()
    print(f"  wrote {arch}")
    research = write_research_sample()
    print(f"  wrote {research}")
    print("\nDone. Next steps:")
    print("  1. python -m auto_applier --cli doctor")
    print("  2. Launch the GUI wizard to parse testpilot.docx into a profile")
    print("  3. python -m auto_applier --cli run --dry-run --limit 1")


if __name__ == "__main__":
    main()
